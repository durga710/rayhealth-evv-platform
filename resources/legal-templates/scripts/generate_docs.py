#!/usr/bin/env python3
"""
generate_docs.py — Build all 25 master .docx templates for the legal-templates package.

Run from resources/legal-templates/:
    /tmp/legal-templates-venv/bin/python scripts/generate_docs.py

Outputs:
    client-onboarding/docx/01..10.docx
    hiring/docx/01..13.docx
    state-addenda/docx/{pennsylvania,ohio}-addendum.docx

Every doc carries:
    - Title heading (24pt bold)
    - Section headings (14pt bold)
    - Body text (11pt)
    - Page numbers in the footer
    - "NOT LEGAL ADVICE" disclaimer footer
    - {{PLACEHOLDER}} markers populated by customize.py from agency-info.json
    - Per-instance blanks (____________________) for fields filled at signing
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
CLIENT_DIR = ROOT / "client-onboarding" / "docx"
HIRING_DIR = ROOT / "hiring" / "docx"
ADDENDA_DIR = ROOT / "state-addenda" / "docx"

for d in (CLIENT_DIR, HIRING_DIR, ADDENDA_DIR):
    d.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

BLANK = "____________________"


def _set_run(run, *, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run(r, size=24, bold=True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run(r, size=12, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_after = Pt(12)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, size=14, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_p(doc: Document, text: str, *, bold=False, italic=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, size=11, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)


def add_callout(doc: Document, text: str, *, size=14) -> None:
    """Bold all-caps regulatory headers (used by HIPAA NPP)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run(r, size=size, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        _set_run(r, size=11)


def add_numbered(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        _set_run(r, size=11)


def add_field_line(doc: Document, label: str) -> None:
    """One-line labeled blank, e.g. 'Client name: ____________________'."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    _set_run(r1, size=11, bold=True)
    r2 = p.add_run(BLANK)
    _set_run(r2, size=11)
    p.paragraph_format.space_after = Pt(4)


def add_signature_block(doc: Document, party_label: str) -> None:
    """Standard signature/printed-name/date block."""
    add_p(doc, "")
    add_field_line(doc, f"{party_label} signature")
    add_field_line(doc, f"{party_label} printed name")
    add_field_line(doc, "Date")


def add_two_party_signatures(doc: Document, left: str, right: str) -> None:
    add_p(doc, "")
    add_h2(doc, "Signatures")
    add_signature_block(doc, left)
    add_signature_block(doc, right)


def add_witness_block(doc: Document) -> None:
    add_field_line(doc, "Witness signature")
    add_field_line(doc, "Witness printed name")
    add_field_line(doc, "Date")


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_disclaimer_footer(doc: Document) -> None:
    """Footer: NOT-LEGAL-ADVICE disclaimer + page number."""
    section = doc.sections[0]
    footer = section.footer

    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(
        "This template is provided as a starting point and does not constitute legal advice. "
        "Have a healthcare-experienced attorney licensed in {{AGENCY_STATE}} review and customize before use."
    )
    _set_run(r1, size=8, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Page ")
    _set_run(r2, size=8, color=RGBColor(0x55, 0x55, 0x55))
    _add_page_number_field(p2)


def add_letterhead(doc: Document) -> None:
    """Top-of-doc agency letterhead block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("{{AGENCY_NAME}}")
    _set_run(r, size=16, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "{{AGENCY_ADDRESS_LINE_1}}, {{AGENCY_ADDRESS_LINE_2}}, "
        "{{AGENCY_CITY}}, {{AGENCY_STATE}} {{AGENCY_ZIP}}"
    )
    _set_run(r2, size=10, color=RGBColor(0x33, 0x33, 0x33))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Phone: {{AGENCY_PHONE}}  |  Fax: {{AGENCY_FAX}}  |  Email: {{AGENCY_EMAIL}}  |  Web: {{AGENCY_WEBSITE}}"
    )
    _set_run(r3, size=10, color=RGBColor(0x33, 0x33, 0x33))
    p3.paragraph_format.space_after = Pt(12)


def new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def save(doc: Document, path: Path) -> Path:
    add_disclaimer_footer(doc)
    doc.save(str(path))
    return path


# ===========================================================================
# CLIENT ONBOARDING (10)
# ===========================================================================


def doc_01_homecare_services_agreement() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Homecare Services Agreement")
    add_subtitle(d, "Effective {{EFFECTIVE_DATE}}")

    add_p(
        d,
        "This Homecare Services Agreement (this \"Agreement\") is entered into between "
        "{{AGENCY_NAME}}, a {{AGENCY_STATE}} {{AGENCY_LEGAL_ENTITY_TYPE}} (the \"Agency\"), "
        "and the client identified below (\"Client\"), on the Effective Date set forth above.",
    )

    add_h2(d, "1. Parties")
    add_field_line(d, "Client legal name")
    add_field_line(d, "Client date of birth")
    add_field_line(d, "Client home address (where services will be provided)")
    add_field_line(d, "Authorized representative (if any)")
    add_field_line(d, "Relationship to Client")

    add_h2(d, "2. Scope of Services")
    add_p(
        d,
        "The Agency will provide non-medical and/or skilled in-home services as authorized by "
        "Client's plan of care, which is incorporated by reference. Services may include personal "
        "care, homemaker services, companion services, respite, and other services within the "
        "Agency's licensure. Services excluded from this Agreement (for example, prescribing or "
        "administering medication outside scope of practice) will be documented in writing.",
    )
    add_field_line(d, "Authorized services (summary)")
    add_field_line(d, "Care plan reference / ID")

    add_h2(d, "3. Schedule and Minimum Visit")
    add_p(
        d,
        "Visits will be scheduled in advance and confirmed by Electronic Visit Verification (EVV) "
        "as required under 21st Century Cures Act, 42 U.S.C. § 1396b(l). Minimum visit length "
        "is {{MINIMUM_VISIT_HOURS}} hour(s). Cancellations made fewer than 24 hours in advance "
        "may be billed at the full scheduled rate, except where prohibited by payer contract or "
        "law.",
    )

    add_h2(d, "4. Rates and Payment")
    add_bullets(
        d,
        [
            "Hourly rate: {{HOURLY_RATE}}",
            "Overtime rate: {{OVERTIME_RATE}}",
            "Holiday rate: {{HOLIDAY_RATE}}",
            "Mileage (where applicable): {{MILEAGE_RATE}} per mile",
            "Invoicing frequency: {{INVOICE_FREQUENCY}}",
            "Payment due: within {{PAYMENT_DUE_DAYS}} days of invoice",
            "Late payment fee: {{LATE_PAYMENT_FEE_PCT}}% per month on unpaid balances",
        ],
    )

    add_h2(d, "5. Electronic Visit Verification (EVV)")
    add_p(
        d,
        "Client acknowledges that caregivers will record clock-in and clock-out at the home using "
        "an Agency-supplied mobile device or alternative EVV method. EVV records may include GPS "
        "location at clock-in/clock-out, time, date, caregiver identity, and tasks performed. "
        "These records are maintained by the Agency and reported to Client's payer where required.",
    )

    add_h2(d, "6. Caregiver Assignment and Substitution")
    add_p(
        d,
        "The Agency selects, employs, and supervises caregivers. The Agency may substitute "
        "caregivers as needed for continuity of care. Client is not the employer of any caregiver "
        "and may not direct, hire, or pay caregivers outside this Agreement.",
    )

    add_h2(d, "7. Term and Termination")
    add_p(
        d,
        "This Agreement begins on the Effective Date and continues until terminated. Either party "
        "may terminate without cause on {{NOTICE_PERIOD_DAYS}} days' written notice. The Agency "
        "may terminate immediately for cause, including unsafe conditions, abuse of caregivers, "
        "non-payment, or material breach. The Agency will, where feasible, assist with continuity "
        "of care upon termination.",
    )

    add_h2(d, "8. Limitation of Liability")
    add_p(
        d,
        "To the maximum extent permitted by law, the Agency's total liability under this "
        "Agreement is limited to the amounts paid by or on behalf of Client to the Agency in the "
        "six (6) months preceding the event giving rise to the claim. Neither party is liable for "
        "indirect, incidental, special, consequential, or punitive damages. Nothing in this "
        "section limits liability that cannot be limited under applicable law (including "
        "liability for gross negligence, willful misconduct, fraud, or personal injury caused by "
        "negligence).",
    )

    add_h2(d, "9. Indemnification")
    add_p(
        d,
        "Each party will indemnify and hold the other harmless from third-party claims arising "
        "out of the indemnifying party's negligence, willful misconduct, or material breach of "
        "this Agreement, subject to the limitation in Section 8.",
    )

    add_h2(d, "10. Insurance")
    add_p(
        d,
        "The Agency maintains general liability insurance with limits of "
        "{{INSURANCE_LIABILITY_LIMIT}}, non-owned automobile coverage of {{INSURANCE_AUTO_LIMIT}}, "
        "and workers' compensation insurance through {{WORKERS_COMP_CARRIER}}. Certificates of "
        "insurance are available on request.",
    )

    add_h2(d, "11. Confidentiality and HIPAA")
    add_p(
        d,
        "The Agency is a covered entity under HIPAA and treats Client's protected health "
        "information (PHI) under its Notice of Privacy Practices, provided separately to Client.",
    )

    add_h2(d, "12. Governing Law and Venue")
    add_p(
        d,
        "This Agreement is governed by the laws of the Commonwealth/State of {{AGENCY_STATE}}, "
        "without regard to its conflict-of-laws principles. Any action arising from this "
        "Agreement will be brought in the state or federal courts located in {{AGENCY_CITY}}, "
        "{{AGENCY_STATE}}.",
    )

    add_h2(d, "13. Entire Agreement")
    add_p(
        d,
        "This Agreement, together with the plan of care, the Notice of Privacy Practices, and "
        "any state addendum, constitutes the entire agreement between the parties on this "
        "subject. Modifications must be in writing and signed by both parties.",
    )

    add_two_party_signatures(d, "Client / authorized representative", "Agency representative")
    return save(d, CLIENT_DIR / "01-homecare-services-agreement.docx")


def doc_02_hipaa_notice_of_privacy_practices() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Notice of Privacy Practices")
    add_subtitle(d, "Effective {{EFFECTIVE_DATE}}  •  Required by 45 CFR § 164.520")

    # Regulator-required header in 14pt or larger, all caps, verbatim.
    add_callout(
        d,
        "THIS NOTICE DESCRIBES HOW MEDICAL INFORMATION ABOUT YOU MAY BE USED AND DISCLOSED AND "
        "HOW YOU CAN GET ACCESS TO THIS INFORMATION. PLEASE REVIEW IT CAREFULLY.",
        size=14,
    )

    add_h2(d, "Who We Are")
    add_p(
        d,
        "{{AGENCY_NAME}} (the \"Agency\") is a HIPAA covered entity. We are required by federal "
        "and state law to protect the privacy of your protected health information (\"PHI\"), "
        "provide you with this Notice describing our legal duties and privacy practices regarding "
        "PHI, and follow the terms of the Notice currently in effect.",
    )

    add_h2(d, "How We May Use and Disclose PHI")
    add_p(d, "Without your written authorization, we may use or disclose your PHI for:")
    add_p(d, "Treatment.", bold=True)
    add_p(
        d,
        "We may use and share your PHI to provide and coordinate your home care, including with "
        "caregivers, supervising nurses, physicians, and other treating providers.",
    )
    add_p(d, "Payment.", bold=True)
    add_p(
        d,
        "We may use and share your PHI to obtain payment for the services we provide, including "
        "with health plans, Medicaid waiver programs, and other payers.",
    )
    add_p(d, "Health care operations.", bold=True)
    add_p(
        d,
        "We may use and share your PHI for our internal operations, including quality improvement, "
        "training, audit, accreditation, and business management.",
    )

    add_h2(d, "Other Permitted Uses and Disclosures (Without Authorization)")
    add_bullets(
        d,
        [
            "When required by law (subpoenas, mandatory reporting, public health authorities).",
            "For public health activities, including reporting communicable disease.",
            "Reports of abuse, neglect, or exploitation to Adult Protective Services or law enforcement, as required by state law.",
            "Health oversight activities such as inspections, audits, and licensure investigations.",
            "Judicial and administrative proceedings, in response to a court order or valid subpoena.",
            "Law enforcement, in limited circumstances permitted by law.",
            "Coroners, medical examiners, and funeral directors.",
            "Organ and tissue donation, if applicable.",
            "Research that has been approved through a privacy-protective process.",
            "To avert a serious threat to health or safety.",
            "Workers' compensation, as authorized by law.",
            "To family members or others involved in your care, only as you have agreed or as we judge to be in your best interest if you are unable to agree.",
        ],
    )

    add_h2(d, "Uses and Disclosures That Require Your Written Authorization")
    add_p(
        d,
        "We will obtain your written authorization before any use or disclosure of PHI for "
        "marketing (other than face-to-face communication or a promotional gift of nominal value), "
        "for any sale of PHI, and for most uses and disclosures of psychotherapy notes. You may "
        "revoke an authorization at any time, in writing, except to the extent we have already "
        "relied on it.",
    )

    add_h2(d, "Your Rights")
    add_bullets(
        d,
        [
            "Right to access: You may inspect and obtain a copy of your PHI in our designated record set, in paper or electronic form. We may charge a reasonable, cost-based fee.",
            "Right to amend: You may request that we amend PHI you believe is incorrect or incomplete. We may deny the request in limited circumstances and will explain why in writing.",
            "Right to an accounting of disclosures: You may request a list of certain disclosures we have made of your PHI in the previous six (6) years.",
            "Right to request restrictions: You may request that we restrict how we use or disclose your PHI for treatment, payment, or operations. We are not generally required to agree, except that we must agree to restrict disclosure to a health plan for services you paid for in full out of pocket.",
            "Right to confidential communications: You may request that we communicate with you by alternative means or at alternative locations.",
            "Right to a paper copy of this Notice, even if you previously agreed to receive it electronically.",
            "Right to be notified of a breach of unsecured PHI affecting you.",
        ],
    )

    add_h2(d, "Our Duties")
    add_bullets(
        d,
        [
            "We are required by law to maintain the privacy and security of your PHI.",
            "We will let you know promptly if a breach occurs that may have compromised the privacy or security of your PHI.",
            "We must follow the duties and privacy practices described in this Notice and give you a copy of it.",
            "We will not use or share your PHI other than as described here unless you tell us we can in writing. If you tell us we can, you may change your mind at any time.",
        ],
    )

    add_h2(d, "Changes to This Notice")
    add_p(
        d,
        "We may change this Notice and apply the new terms to PHI we already have, as well as PHI "
        "we receive in the future. We will post the current Notice in our office and on "
        "{{AGENCY_WEBSITE}}, and we will provide a copy on request.",
    )

    add_h2(d, "Complaints")
    add_p(
        d,
        "If you believe your privacy rights have been violated, you may complain to us or to the "
        "U.S. Department of Health and Human Services, Office for Civil Rights. We will not "
        "retaliate against you for filing a complaint.",
    )
    add_bullets(
        d,
        [
            "Agency Privacy Officer: {{PRIVACY_OFFICER_NAME}}, {{PRIVACY_OFFICER_PHONE}}, {{PRIVACY_OFFICER_EMAIL}}",
            "U.S. Department of Health and Human Services, Office for Civil Rights, 200 Independence Avenue, S.W., Washington, D.C. 20201",
            "OCR phone: 1-877-696-6775  •  OCR online complaint: https://www.hhs.gov/ocr/complaints",
        ],
    )

    add_h2(d, "Acknowledgement of Receipt")
    add_p(
        d,
        "I acknowledge that I have received a copy of {{AGENCY_NAME}}'s Notice of Privacy "
        "Practices.",
    )
    add_signature_block(d, "Client / personal representative")

    return save(d, CLIENT_DIR / "02-hipaa-notice-of-privacy-practices.docx")


def doc_03_hipaa_authorization() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "HIPAA Authorization for Use or Disclosure of PHI")
    add_subtitle(d, "Per 45 CFR § 164.508")

    add_h2(d, "Client / Patient")
    add_field_line(d, "Full legal name")
    add_field_line(d, "Date of birth")
    add_field_line(d, "Address")

    add_h2(d, "1. Information to Be Used or Disclosed (Specific Description)")
    add_p(
        d,
        "Describe the specific PHI to be used or disclosed (for example: \"plan of care, visit "
        "notes, and medication list dated [range]\"). Be specific — \"all medical records\" "
        "is broader than usually appropriate.",
    )
    add_field_line(d, "Specific PHI to be released")

    add_h2(d, "2. Person(s) or Class of Persons Authorized to Make the Disclosure")
    add_p(d, "{{AGENCY_NAME}} and its workforce members.")

    add_h2(d, "3. Person(s) or Class of Persons Authorized to Receive the Disclosure")
    add_field_line(d, "Recipient name / organization")
    add_field_line(d, "Recipient address")
    add_field_line(d, "Recipient phone")

    add_h2(d, "4. Purpose of the Use or Disclosure")
    add_p(
        d,
        "If you do not wish to state a purpose, you may write \"at the request of the individual.\"",
    )
    add_field_line(d, "Purpose")

    add_h2(d, "5. Expiration")
    add_p(d, "This authorization expires on the earliest of:")
    add_bullets(
        d,
        [
            "The date or event written here: ____________________",
            "One (1) year from the date signed below if no other expiration is given.",
            "Written revocation by the individual or the individual's personal representative.",
        ],
    )

    add_h2(d, "6. Right to Revoke")
    add_p(
        d,
        "You have the right to revoke this authorization in writing at any time by sending a "
        "signed letter to the Agency Privacy Officer at {{PRIVACY_OFFICER_EMAIL}} or "
        "{{AGENCY_ADDRESS_LINE_1}}, {{AGENCY_CITY}}, {{AGENCY_STATE}} {{AGENCY_ZIP}}. Revocation "
        "will not apply to information we have already used or disclosed in reliance on this "
        "authorization.",
    )

    add_h2(d, "7. Conditioning of Treatment")
    add_p(
        d,
        "We cannot condition your treatment, payment, enrollment in a health plan, or eligibility "
        "for benefits on whether you sign this authorization, except in limited circumstances "
        "permitted by 45 CFR § 164.508(b)(4).",
    )

    add_h2(d, "8. Potential for Redisclosure")
    add_p(
        d,
        "Information used or disclosed under this authorization may be redisclosed by the "
        "recipient and no longer protected by federal privacy regulations. Special rules may "
        "apply to substance use disorder records (42 CFR Part 2), HIV-related information, "
        "mental health records, and genetic information.",
    )

    add_h2(d, "9. Signature")
    add_p(
        d,
        "I have read this authorization and understand my rights. I am signing it voluntarily.",
    )
    add_signature_block(d, "Client / personal representative")
    add_field_line(d, "If signed by a personal representative, describe authority")

    return save(d, CLIENT_DIR / "03-hipaa-authorization-for-use-disclosure.docx")


def doc_04_client_bill_of_rights() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Client Bill of Rights and Responsibilities")

    add_p(
        d,
        "{{AGENCY_NAME}} respects your dignity, privacy, and right to make informed decisions "
        "about your care. As a client of the Agency, you have the following rights and "
        "responsibilities. These rights are protected by 28 Pa. Code Ch. 611 (in Pennsylvania), "
        "ORC Ch. 3740 / OAC 3701-13 (in Ohio), and other applicable law.",
    )

    add_h2(d, "Your Rights")
    add_numbered(
        d,
        [
            "To be treated with courtesy, respect, and recognition of your dignity and individuality.",
            "To receive care without discrimination on the basis of race, color, national origin, sex, age, religion, disability, sexual orientation, gender identity, source of payment, or any other basis prohibited by law.",
            "To be informed, in advance, of the services to be provided, the names and qualifications of the people who will provide your care, and the rates you will be charged.",
            "To participate in the development and revision of your plan of care.",
            "To be informed, in advance, of any change in your plan of care.",
            "To refuse care, treatment, or services after being informed of the consequences of refusal.",
            "To have your property and home treated with respect.",
            "To privacy and confidentiality of your records, in accordance with HIPAA and state law.",
            "To voice grievances about care or services without fear of retaliation.",
            "To be informed of the procedure for filing a grievance with the Agency, with the state regulator ({{STATE_REGULATOR_NAME}} at {{STATE_REGULATOR_PHONE}}), and with Adult Protective Services ({{APS_PHONE}}).",
            "To be free from physical and mental abuse, neglect, exploitation, and misappropriation of your property.",
            "To formulate advance directives and have those directives respected.",
            "To be informed of the Agency's ownership and any controlling interests.",
            "To be informed in writing of any limitations on the Agency's ability to provide care.",
        ],
    )

    add_h2(d, "Your Responsibilities")
    add_numbered(
        d,
        [
            "Provide accurate and complete information about your medical history, current condition, medications, and other matters relevant to your care.",
            "Participate in your plan of care and inform the Agency of changes in your condition, environment, or contact information.",
            "Provide a safe environment for caregivers, including controlling pets, removing hazards, and refraining from smoking during visits if requested.",
            "Treat caregivers with respect and refrain from discrimination, harassment, or violence.",
            "Pay for services on the agreed schedule, or make timely arrangements with the Agency if payment will be delayed.",
            "Provide reasonable notice if you must cancel or reschedule a visit.",
            "Inform the Agency promptly of any concerns, complaints, or changes in payer authorization.",
        ],
    )

    add_h2(d, "Acknowledgement")
    add_p(d, "I have received and reviewed this Bill of Rights and Responsibilities.")
    add_signature_block(d, "Client / personal representative")

    return save(d, CLIENT_DIR / "04-client-bill-of-rights-and-responsibilities.docx")


def doc_05_consent_to_care() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Consent to In-Home Care")

    add_h2(d, "1. Voluntary Consent")
    add_p(
        d,
        "I voluntarily consent to receive in-home care services from {{AGENCY_NAME}}. I "
        "understand that I am free to refuse any service, treatment, or task, and to discontinue "
        "services at any time, subject to the notice provisions of my services agreement.",
    )

    add_h2(d, "2. Care to Be Provided")
    add_p(
        d,
        "Services will be performed in accordance with my plan of care, which I have had the "
        "opportunity to review and discuss. I understand that the plan of care may be revised "
        "based on my needs, and that I will be informed of significant changes.",
    )

    add_h2(d, "3. Risks and Benefits")
    add_p(
        d,
        "I understand that no medical service is risk-free. I have had the opportunity to ask "
        "questions about the benefits, risks, and alternatives to home care.",
    )

    add_h2(d, "4. Caregivers")
    add_p(
        d,
        "I consent to having caregivers employed or supervised by the Agency enter my home to "
        "perform authorized services. I understand caregivers may rotate based on availability "
        "and continuity of care.",
    )

    add_h2(d, "5. Emergency Care")
    add_p(
        d,
        "If a medical emergency arises during a visit, I authorize Agency personnel to call 911, "
        "request emergency medical services, and provide emergency information to first "
        "responders. The Agency will notify my emergency contact as soon as practical.",
    )

    add_h2(d, "6. Photography for Care Coordination")
    add_p(
        d,
        "I understand that caregivers may take limited photographs of skin conditions, wounds, "
        "or environmental hazards solely for clinical documentation. These photographs are "
        "treated as PHI. Photography for any other purpose requires a separate Photo / Media "
        "Release.",
    )

    add_h2(d, "7. Right to Withdraw Consent")
    add_p(
        d,
        "I may withdraw this consent at any time by notifying the Agency in writing. Withdrawal "
        "does not affect care already provided.",
    )

    add_signature_block(d, "Client / personal representative")
    add_witness_block(d)

    return save(d, CLIENT_DIR / "05-consent-to-care.docx")


def doc_06_financial_responsibility() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Financial Responsibility & Payment Authorization")

    add_h2(d, "1. Rate Schedule")
    add_bullets(
        d,
        [
            "Hourly rate: {{HOURLY_RATE}}",
            "Overtime rate: {{OVERTIME_RATE}}",
            "Holiday rate: {{HOLIDAY_RATE}}",
            "Mileage (where applicable): {{MILEAGE_RATE}} per mile",
            "Minimum visit length: {{MINIMUM_VISIT_HOURS}} hour(s)",
        ],
    )

    add_h2(d, "2. Billing")
    add_p(
        d,
        "Invoices are issued on a {{INVOICE_FREQUENCY}} basis and are due within "
        "{{PAYMENT_DUE_DAYS}} days of issuance. Past-due balances accrue a service charge of "
        "{{LATE_PAYMENT_FEE_PCT}}% per month or the maximum allowed by law, whichever is less.",
    )

    add_h2(d, "3. Payer of Record")
    add_field_line(d, "Primary payer (private pay / Medicaid / waiver / VA / LTC insurance / other)")
    add_field_line(d, "Payer policy or member ID")
    add_field_line(d, "Authorization / case number")
    add_field_line(d, "Authorization period")

    add_h2(d, "4. Assignment of Benefits")
    add_p(
        d,
        "I assign to {{AGENCY_NAME}} all rights, benefits, and authorized payments that are or "
        "may become due from any insurance company, Medicaid, Medicare, waiver program, or other "
        "payer for services rendered by the Agency. I authorize the Agency to release "
        "information needed to process claims.",
    )

    add_h2(d, "5. Client Financial Responsibility")
    add_p(
        d,
        "I understand that I am ultimately responsible for the cost of services not covered by my "
        "payer, including private-pay hours, deductibles, copayments, and coinsurance. The "
        "Agency will provide a written estimate on request.",
    )

    add_h2(d, "6. Payment Method")
    add_field_line(d, "Method (check / ACH / credit card / payroll deduction / other)")
    add_field_line(d, "Account or card last four")
    add_field_line(d, "Billing address (if different from home address)")

    add_h2(d, "7. Disputes")
    add_p(
        d,
        "I will notify the Agency in writing within thirty (30) days of receipt of any invoice I "
        "wish to dispute. The Agency will respond within fifteen (15) business days.",
    )

    add_signature_block(d, "Client / personal representative")
    add_witness_block(d)

    return save(d, CLIENT_DIR / "06-financial-responsibility-and-payment-authorization.docx")


def doc_07_advance_directives() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Advance Directives Acknowledgment")
    add_subtitle(d, "Patient Self-Determination Act, 42 U.S.C. § 1395cc(f)")

    add_p(
        d,
        "Federal law gives you the right to make decisions about your medical care, including "
        "the right to accept or refuse treatment and the right to formulate an advance directive. "
        "An advance directive is a written instruction (such as a living will, durable power of "
        "attorney for health care, or POLST/MOLST/POST form) that takes effect if you become "
        "unable to make decisions for yourself.",
    )

    add_h2(d, "1. Agency Policy")
    add_p(
        d,
        "{{AGENCY_NAME}} will honor advance directives that are valid under the law of "
        "{{AGENCY_STATE}}. The Agency itself does not provide notarial or witness services for "
        "executing advance directives, but caregivers and staff will refer you to qualified "
        "resources on request.",
    )

    add_h2(d, "2. Your Status")
    add_bullets(
        d,
        [
            "[ ] I have an advance directive. A copy is on file with the Agency.",
            "[ ] I have an advance directive. A copy is NOT on file; I will provide one within fourteen (14) days.",
            "[ ] I do not have an advance directive and do not wish to have one at this time.",
            "[ ] I would like more information about advance directives.",
        ],
    )

    add_h2(d, "3. Health Care Surrogate")
    add_field_line(d, "Surrogate / health-care agent name")
    add_field_line(d, "Surrogate phone")
    add_field_line(d, "Surrogate relationship")

    add_h2(d, "4. POLST / MOLST / POST")
    add_field_line(d, "Do you have a POLST/MOLST/POST? (yes / no)")
    add_field_line(d, "Code status (full code / DNR / other)")
    add_p(
        d,
        "Caregivers will follow your written code status as documented in the plan of care and "
        "any POLST/MOLST/POST form on file. In the absence of written documentation, caregivers "
        "will initiate emergency response and call 911.",
    )

    add_h2(d, "5. Acknowledgment")
    add_p(
        d,
        "I have received written information from the Agency describing my rights under "
        "{{AGENCY_STATE}} law to make decisions concerning medical care, including the right to "
        "accept or refuse treatment and the right to formulate advance directives.",
    )
    add_signature_block(d, "Client / personal representative")

    return save(d, CLIENT_DIR / "07-advance-directives-acknowledgment.docx")


def doc_08_photo_media_release() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Photo & Media Release (Optional)")

    add_p(
        d,
        "This release is OPTIONAL. You are not required to sign it as a condition of receiving "
        "services. Signing this release allows {{AGENCY_NAME}} to use your image and likeness in "
        "the limited ways you authorize below. You may revoke this release at any time in "
        "writing by contacting {{PRIVACY_OFFICER_EMAIL}}, with no effect on services.",
    )

    add_h2(d, "1. Permitted Uses (check each that you authorize)")
    add_bullets(
        d,
        [
            "[ ] Internal newsletters and bulletin boards (not distributed outside the Agency).",
            "[ ] Marketing materials (brochures, flyers, print advertising).",
            "[ ] Agency website and social media (Facebook, Instagram, LinkedIn).",
            "[ ] Local news media or press releases.",
            "[ ] Training and education for caregivers.",
        ],
    )
    add_p(
        d,
        "Uses NOT authorized include any use other than those checked above. Photographs taken "
        "for clinical documentation (wounds, environmental hazards) are NOT covered by this "
        "release; those are handled as PHI under HIPAA.",
    )

    add_h2(d, "2. Description of Media")
    add_field_line(d, "Type of media (photo / video / audio / quote)")
    add_field_line(d, "Date taken")
    add_field_line(d, "Brief description")

    add_h2(d, "3. Compensation")
    add_p(
        d,
        "I understand I will not receive monetary compensation for the use of my image or "
        "likeness under this release.",
    )

    add_h2(d, "4. Release of Claims")
    add_p(
        d,
        "I release the Agency, its employees, and authorized partners from claims arising out of "
        "the authorized uses checked above, except for claims arising from negligence, willful "
        "misconduct, or use beyond what I authorized.",
    )

    add_h2(d, "5. Revocation")
    add_p(
        d,
        "I may revoke this release at any time by sending written notice to "
        "{{PRIVACY_OFFICER_EMAIL}}. Revocation will not require recall of materials already "
        "printed or distributed before revocation, but the Agency will stop further use within a "
        "reasonable period.",
    )

    add_signature_block(d, "Client / personal representative")
    add_witness_block(d)

    return save(d, CLIENT_DIR / "08-photo-media-release.docx")


def doc_09_emergency_contact() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Emergency Contact & Critical Information")

    add_h2(d, "Client")
    add_field_line(d, "Full legal name")
    add_field_line(d, "Date of birth")
    add_field_line(d, "Home address")
    add_field_line(d, "Best phone number")
    add_field_line(d, "Preferred language")

    add_h2(d, "Emergency Contacts (in order)")
    for i in (1, 2, 3):
        add_field_line(d, f"#{i} Name")
        add_field_line(d, f"#{i} Relationship")
        add_field_line(d, f"#{i} Phone")
        add_field_line(d, f"#{i} Has authority to make medical decisions? (yes / no)")
        add_p(d, "")

    add_h2(d, "Allergies and Sensitivities")
    add_field_line(d, "Drug allergies")
    add_field_line(d, "Food / latex / environmental allergies")
    add_field_line(d, "Reaction(s)")

    add_h2(d, "Active Medications")
    add_p(d, "Attach a current medication list, or use the lines below.")
    for _ in range(6):
        add_field_line(d, "Medication / dose / frequency / prescriber")

    add_h2(d, "Code Status / Advance Directive")
    add_field_line(d, "Code status (full code / DNR / other)")
    add_field_line(d, "Living will on file? (yes / no)")
    add_field_line(d, "Health care power of attorney? (yes / no)")
    add_field_line(d, "POLST / MOLST / POST? (yes / no)")

    add_h2(d, "Hospital Preference")
    add_field_line(d, "Preferred hospital")
    add_field_line(d, "Preferred ER if different")

    add_h2(d, "Primary Care and Specialty Providers")
    add_field_line(d, "Primary care provider name / phone")
    add_field_line(d, "Pharmacy name / phone")
    add_field_line(d, "Specialist 1 name / phone")
    add_field_line(d, "Specialist 2 name / phone")

    add_h2(d, "Other Critical Information")
    add_field_line(d, "Mobility / fall risk")
    add_field_line(d, "Communication aids (hearing aid / glasses / dentures)")
    add_field_line(d, "Behavioral / cognitive considerations")
    add_field_line(d, "Pets in the home")
    add_field_line(d, "Firearms in the home? (yes / no)")

    add_signature_block(d, "Client / personal representative")
    add_field_line(d, "Information collected by (caregiver / nurse)")
    add_field_line(d, "Date collected")

    return save(d, CLIENT_DIR / "09-emergency-contact-and-critical-information.docx")


def doc_10_grievance() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Grievance and Complaint Policy")

    add_p(
        d,
        "{{AGENCY_NAME}} encourages clients, families, and caregivers to raise concerns. We "
        "investigate every grievance, take it seriously, and will not retaliate against anyone "
        "for filing one in good faith.",
    )

    add_h2(d, "1. How to File a Grievance Internally")
    add_numbered(
        d,
        [
            "Speak to your assigned caregiver or supervisor first if you feel comfortable doing so.",
            "If unresolved, contact the Compliance Officer: {{COMPLIANCE_OFFICER_NAME}}, {{COMPLIANCE_OFFICER_EMAIL}}.",
            "Submit your grievance in writing if possible — by email, mail to {{AGENCY_ADDRESS_LINE_1}}, {{AGENCY_CITY}}, {{AGENCY_STATE}} {{AGENCY_ZIP}}, or in person.",
            "We will acknowledge receipt within five (5) business days.",
            "We will investigate and respond in writing within thirty (30) calendar days, with a longer timeline only if the matter is unusually complex (with notice).",
        ],
    )

    add_h2(d, "2. External Resources")
    add_bullets(
        d,
        [
            "State home care regulator: {{STATE_REGULATOR_NAME}}, {{STATE_REGULATOR_PHONE}}, {{STATE_REGULATOR_WEBSITE}}",
            "Adult Protective Services: {{APS_PHONE}}, {{APS_WEBSITE}}",
            "Long-Term Care Ombudsman (state-specific). Caregivers can provide the local number on request.",
            "Office for Civil Rights (HIPAA complaints): 1-877-696-6775, https://www.hhs.gov/ocr/complaints",
            "If you believe there is an immediate safety concern: call 911.",
        ],
    )

    add_h2(d, "3. Confidentiality and Non-Retaliation")
    add_p(
        d,
        "Your grievance will be handled as confidentially as practical. The Agency prohibits "
        "retaliation against any client, family member, or employee for filing a grievance or "
        "cooperating in an investigation, in accordance with applicable law.",
    )

    add_h2(d, "4. Acknowledgement")
    add_signature_block(d, "Client / personal representative")
    return save(d, CLIENT_DIR / "10-grievance-and-complaint-policy.docx")


# ===========================================================================
# HIRING / HR (13)
# ===========================================================================


def doc_h01_employment_application() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Employment Application")

    add_p(
        d,
        "{{AGENCY_NAME}} is an Equal Opportunity Employer. We do not discriminate on the basis of "
        "race, color, religion, national origin, sex, sexual orientation, gender identity, age, "
        "disability, genetic information, veteran status, or any other status protected by "
        "federal, state, or local law. We comply with the Americans with Disabilities Act and "
        "will provide reasonable accommodations to applicants and employees on request.",
    )

    add_h2(d, "1. Personal Information")
    add_field_line(d, "Full legal name")
    add_field_line(d, "Phone")
    add_field_line(d, "Email")
    add_field_line(d, "Mailing address")
    add_field_line(d, "Position applied for")
    add_field_line(d, "How did you hear about us?")
    add_field_line(d, "Are you legally authorized to work in the United States? (yes / no)")
    add_field_line(d, "Will you require sponsorship to work in the United States? (yes / no)")

    add_h2(d, "2. Availability")
    add_field_line(d, "Date available to start")
    add_field_line(d, "Desired hours per week")
    add_field_line(d, "Days/shifts available")
    add_field_line(d, "Do you have reliable transportation? (yes / no)")

    add_h2(d, "3. Education")
    for label in ("High school / GED", "College or trade school", "Other relevant training"):
        add_field_line(d, f"{label} - institution / years / completed (yes / no)")

    add_h2(d, "4. Licenses and Certifications")
    add_p(
        d,
        "List all current licenses, certifications, and training relevant to home care (e.g., CNA, "
        "HHA, CPR/BLS, First Aid, dementia care, medication aide).",
    )
    for _ in range(4):
        add_field_line(d, "Credential / issuing body / number / expiration")

    add_h2(d, "5. Employment History (most recent first)")
    for i in (1, 2, 3):
        add_field_line(d, f"#{i} Employer / job title")
        add_field_line(d, f"#{i} Dates of employment")
        add_field_line(d, f"#{i} Supervisor and contact")
        add_field_line(d, f"#{i} Reason for leaving")
        add_field_line(d, f"#{i} May we contact this employer? (yes / no)")
        add_p(d, "")

    add_h2(d, "6. References")
    add_p(d, "Provide at least two professional references (not related by blood or marriage).")
    for i in (1, 2, 3):
        add_field_line(d, f"#{i} Name / relationship / phone / email")

    add_h2(d, "7. Background and Eligibility")
    add_p(
        d,
        "We comply with state and local \"ban-the-box\" laws. The Agency requires criminal "
        "background checks for caregiving positions because we serve a vulnerable population. "
        "Answers below will be considered in light of the nature of the offense, time elapsed, "
        "and relevance to the position.",
    )
    add_field_line(d, "Have you ever been on the OIG List of Excluded Individuals/Entities? (yes / no)")
    add_field_line(d, "Are you on any state Medicaid exclusion list? (yes / no)")
    add_field_line(d, "Are you on the {{AGENCY_STATE}} sex offender or abuse registry? (yes / no)")

    add_h2(d, "8. Certification by Applicant")
    add_p(
        d,
        "I certify that the information in this application is true and complete to the best of "
        "my knowledge. I understand that any false statement, omission, or misrepresentation may "
        "result in withdrawal of an offer or termination of employment. I authorize the Agency to "
        "verify any information given here by contacting employers, schools, references, and "
        "credentialing bodies.",
    )
    add_signature_block(d, "Applicant")

    return save(d, HIRING_DIR / "01-employment-application.docx")


def doc_h02_offer_letter() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Offer of Employment")

    add_p(d, "Date: ____________________")
    add_p(d, "To: ____________________  (\"Candidate\")")

    add_p(d, "Dear Candidate,")
    add_p(
        d,
        "On behalf of {{AGENCY_NAME}}, I am pleased to offer you employment in the position "
        "described below. Please review this letter carefully and, if you accept, sign and return "
        "it by the date listed at the bottom.",
    )

    add_h2(d, "1. Position")
    add_field_line(d, "Position title")
    add_field_line(d, "Reports to")
    add_field_line(d, "Classification (FLSA: non-exempt / exempt)")
    add_field_line(d, "Status (full-time / part-time / per-diem)")
    add_field_line(d, "Anticipated start date")

    add_h2(d, "2. Compensation")
    add_p(
        d,
        "Your starting compensation is {{HOURLY_RATE}} per hour, paid {{INVOICE_FREQUENCY}}. "
        "Overtime is paid at {{OVERTIME_RATE}} for hours worked in excess of 40 in a workweek, "
        "consistent with the Fair Labor Standards Act and {{AGENCY_STATE}} law. Holiday hours, "
        "where applicable, are paid at {{HOLIDAY_RATE}}.",
    )

    add_h2(d, "3. Benefits")
    add_p(
        d,
        "You are eligible for the benefits described in the Agency's employee handbook (subject "
        "to the eligibility rules in each plan). Benefit programs may change from time to time.",
    )

    add_h2(d, "4. Contingencies")
    add_p(
        d,
        "This offer is contingent on each of the following, which must be completed before your "
        "first scheduled visit:",
    )
    add_bullets(
        d,
        [
            "Successful completion of a state and federal criminal background check.",
            "Verification of identity and employment authorization (Form I-9).",
            "Negative pre-employment drug screen.",
            "Verification of all licenses, certifications, and credentials listed in your application.",
            "Two-step TB screening (or chest X-ray) where required by state law.",
            "Signed receipt of the Employee Handbook, Code of Conduct, and HIPAA training acknowledgment.",
            "If your role requires driving: a satisfactory motor vehicle record (MVR) and proof of personal auto insurance with at least the limits required by Agency policy.",
        ],
    )

    add_h2(d, "5. At-Will Employment")
    add_p(
        d,
        "Your employment is \"at-will,\" which means that either you or the Agency may end the "
        "employment relationship at any time, with or without cause and with or without notice, "
        "subject to applicable law. No representation by any Agency employee can change this "
        "at-will relationship except a written agreement signed by an Agency officer.",
    )

    add_h2(d, "6. Confidentiality")
    add_p(
        d,
        "By accepting this offer, you agree to sign and abide by the Agency's Confidentiality / "
        "Non-Disclosure Agreement, which protects PHI and other confidential information.",
    )

    add_h2(d, "7. Acceptance")
    add_p(d, "Please indicate acceptance by signing below and returning this letter by ____________________.")
    add_signature_block(d, "Candidate")
    add_signature_block(d, "Agency representative")

    return save(d, HIRING_DIR / "02-offer-letter.docx")


def doc_h03_confidentiality_nda() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Confidentiality & Non-Disclosure Agreement")

    add_p(
        d,
        "This Confidentiality and Non-Disclosure Agreement (this \"Agreement\") is between "
        "{{AGENCY_NAME}} (the \"Agency\") and the individual signing below (\"Employee\"). It is "
        "effective on the date Employee signs.",
    )

    add_h2(d, "1. Confidential Information")
    add_p(
        d,
        "\"Confidential Information\" means all non-public information learned in the course of "
        "employment, including:",
    )
    add_bullets(
        d,
        [
            "Protected Health Information (\"PHI\") of clients, as defined by HIPAA, including names, addresses, diagnoses, schedules, plans of care, and any information that could identify a client.",
            "Personnel information about caregivers and staff.",
            "Business information of the Agency: payer rates, vendor lists, financial data, training materials, internal policies, software credentials.",
            "Information about clients' homes: alarm codes, key locations, security routines, family relationships.",
        ],
    )

    add_h2(d, "2. Use and Non-Disclosure")
    add_p(
        d,
        "Employee will use Confidential Information only as necessary to perform Employee's job. "
        "Employee will not share Confidential Information with anyone inside or outside the "
        "Agency unless that person has a clear need to know and is authorized to receive it.",
    )

    add_h2(d, "3. HIPAA")
    add_p(
        d,
        "Employee will follow all Agency HIPAA policies. Unauthorized access, use, or disclosure "
        "of PHI may result in disciplinary action up to and including termination, civil penalties, "
        "and criminal prosecution under 42 U.S.C. § 1320d-6.",
    )

    add_h2(d, "4. Social Media and Public Statements")
    add_p(
        d,
        "Employee will not post photographs, videos, audio, or written descriptions of clients, "
        "client homes, or co-workers on social media, blogs, podcasts, group chats, or any other "
        "public or semi-public forum without express written permission from the Agency.",
    )

    add_h2(d, "5. Return of Materials")
    add_p(
        d,
        "On termination of employment for any reason, Employee will return all Agency property "
        "and Confidential Information in any form (paper, mobile device, key, badge, EVV device) "
        "and will permanently delete any Agency information from personal devices.",
    )

    add_h2(d, "6. Reporting Obligations Preserved")
    add_p(
        d,
        "Nothing in this Agreement prevents Employee from reporting potential violations of law to "
        "any government agency, participating in a government investigation, or making other "
        "disclosures protected by whistleblower laws. Nothing in this Agreement prevents Employee "
        "from discussing wages, hours, or working conditions as protected by the National Labor "
        "Relations Act.",
    )

    add_h2(d, "7. Survival")
    add_p(d, "Sections 1-6 survive termination of employment.")

    add_h2(d, "8. Governing Law")
    add_p(
        d,
        "This Agreement is governed by the law of {{AGENCY_STATE}}, without regard to its "
        "conflict-of-laws principles.",
    )

    add_signature_block(d, "Employee")
    add_signature_block(d, "Agency representative")

    return save(d, HIRING_DIR / "03-confidentiality-nda.docx")


def doc_h04_background_check_fcra() -> Path:
    """STANDALONE FCRA disclosure + auth — important regulatory requirement."""
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Disclosure Regarding Background Investigation")
    add_subtitle(d, "Standalone disclosure under the Fair Credit Reporting Act, 15 U.S.C. § 1681b(b)(2)(A)")

    add_p(
        d,
        "{{AGENCY_NAME}} (the \"Agency\") may obtain information about you from a consumer "
        "reporting agency for employment purposes. Thus, you may be the subject of a \"consumer "
        "report\" and/or an \"investigative consumer report\" which may include information about "
        "your character, general reputation, personal characteristics, and/or mode of living. "
        "These reports may be obtained at any time after your receipt of this disclosure and, if "
        "you are hired, throughout your employment.",
    )
    add_p(
        d,
        "Reports may include, but are not limited to: criminal history; sex offender registry "
        "checks; abuse, neglect, and exploitation registry checks (including the {{AGENCY_STATE}} "
        "registry); OIG List of Excluded Individuals/Entities; state Medicaid exclusion lists; "
        "education and credential verification; employment verification; reference interviews; "
        "and motor vehicle records.",
    )
    add_p(
        d,
        "The information may be obtained from public records, court records, federal and state "
        "agencies, schools, employers, references, and other sources. You have the right, on "
        "written request to the Agency, to request the nature and scope of any investigative "
        "consumer report.",
    )

    add_p(
        d,
        "This document is provided to you alone and is not combined with any other employment "
        "document. The Authorization that follows is on a separate page.",
        italic=True,
    )

    d.add_page_break()

    add_title(d, "Authorization for Background Investigation")

    add_p(
        d,
        "I have received the Disclosure Regarding Background Investigation and a copy of \"A "
        "Summary of Your Rights Under the Fair Credit Reporting Act.\"",
    )
    add_p(
        d,
        "I authorize {{AGENCY_NAME}} and its background-check vendor to obtain a consumer report "
        "and/or investigative consumer report about me for employment purposes, both before and "
        "(if hired) during my employment.",
    )
    add_p(
        d,
        "I authorize courts, law enforcement agencies, schools, employers, references, and other "
        "custodians of records to release information about me to the Agency or its vendor.",
    )

    add_h2(d, "Identifiers (used only to obtain accurate records)")
    add_field_line(d, "Full legal name (first / middle / last / suffix)")
    add_field_line(d, "Other names used")
    add_field_line(d, "Date of birth (used only for matching, not for hiring decisions)")
    add_field_line(d, "Driver's license number / state (if relevant to position)")

    add_h2(d, "State-Specific Notices")
    add_p(d, "Pennsylvania.", bold=True)
    add_p(
        d,
        "An employer in Pennsylvania may consider felony and misdemeanor convictions only to the "
        "extent they relate to suitability for the position. The Agency complies with the "
        "Pennsylvania Criminal History Record Information Act (CHRIA), 18 Pa. C.S. § 9125, "
        "and the Older Adults Protective Services Act (OAPSA), 35 P.S. § 10225.501 et seq.",
    )
    add_p(d, "Ohio.", bold=True)
    add_p(
        d,
        "Ohio criminal records checks for home health and home care workers are conducted by the "
        "Ohio Bureau of Criminal Investigation under ORC § 109.572 and OAC 3701-13. The "
        "Agency complies with these requirements.",
    )

    add_signature_block(d, "Applicant / Employee")
    return save(d, HIRING_DIR / "04-background-check-authorization-fcra.docx")


def doc_h05_drug_alcohol() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Drug and Alcohol Testing Consent")

    add_p(
        d,
        "{{AGENCY_NAME}} maintains a drug-free workplace. To protect the safety of clients, "
        "caregivers, and the public, the Agency requires drug and alcohol testing under the "
        "circumstances described below. This consent is a condition of employment.",
    )

    add_h2(d, "1. Pre-Employment Testing")
    add_p(
        d,
        "All offers of employment are contingent on a negative pre-employment drug screen. "
        "Specimens are collected by an Agency-approved laboratory using a Substance Abuse and "
        "Mental Health Services Administration (SAMHSA)-certified panel.",
    )

    add_h2(d, "2. Reasonable Suspicion Testing")
    add_p(
        d,
        "Employees may be tested when a trained supervisor has a reasonable, documented suspicion "
        "of impairment based on objective criteria, including observed behavior, speech, body "
        "odors, or physical signs and symptoms.",
    )

    add_h2(d, "3. Post-Incident Testing")
    add_p(
        d,
        "Employees may be tested following a workplace accident or incident that involves "
        "personal injury, property damage, a medication error, a fall while on duty, or a motor "
        "vehicle collision while operating a vehicle for work.",
    )

    add_h2(d, "4. Random Testing (if applicable to position)")
    add_field_line(d, "Position subject to random testing? (yes / no)")
    add_p(
        d,
        "If yes, employees in this position may be selected for unannounced testing using a "
        "neutral, computer-generated process.",
    )

    add_h2(d, "5. Confidentiality and Medical Review Officer (MRO)")
    add_p(
        d,
        "All test results are confidential. Positive results are reviewed by a Medical Review "
        "Officer who will give the employee an opportunity to provide a legitimate medical "
        "explanation (such as a current prescription) before reporting to the Agency.",
    )

    add_h2(d, "6. Consequences")
    add_p(
        d,
        "A confirmed positive test, refusal to test, tampering with a specimen, or substituted "
        "specimen will result in withdrawal of an offer or disciplinary action up to and "
        "including termination, consistent with applicable law.",
    )

    add_h2(d, "7. Authorization")
    add_p(
        d,
        "I consent to drug and alcohol testing as described above. I authorize the testing "
        "facility and the MRO to release results to the Agency for employment purposes only.",
    )
    add_signature_block(d, "Applicant / Employee")

    return save(d, HIRING_DIR / "05-drug-and-alcohol-testing-consent.docx")


def doc_h06_direct_deposit() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Direct Deposit Authorization")

    add_p(
        d,
        "I authorize {{AGENCY_NAME}} to deposit my regular wages into the bank account(s) listed "
        "below. I understand that deposits may take one to two pay cycles to begin and that I am "
        "responsible for confirming the first deposit posts correctly. I will notify the Agency "
        "in writing of any change to account information.",
    )

    add_h2(d, "1. Account #1")
    add_field_line(d, "Bank name")
    add_field_line(d, "Account type (checking / savings)")
    add_field_line(d, "Routing number (9 digits)")
    add_field_line(d, "Account number")
    add_field_line(d, "Allocation (full deposit / specific dollar amount / specific %)")

    add_h2(d, "2. Account #2 (optional)")
    add_field_line(d, "Bank name")
    add_field_line(d, "Account type (checking / savings)")
    add_field_line(d, "Routing number")
    add_field_line(d, "Account number")
    add_field_line(d, "Allocation")

    add_h2(d, "3. Voided Check or Bank Letter Required")
    add_p(
        d,
        "Attach a voided check or a bank letter confirming the routing and account numbers. "
        "Direct deposit cannot be set up without one of these documents.",
    )

    add_h2(d, "4. Reversals")
    add_p(
        d,
        "I authorize the Agency to debit my account to recover any amount erroneously deposited, "
        "subject to applicable banking and wage payment law.",
    )

    add_h2(d, "5. Cancellation")
    add_p(d, "This authorization remains in effect until I revoke it in writing or my employment ends.")

    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "06-direct-deposit-authorization.docx")


def doc_h07_handbook_ack() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Employee Handbook Acknowledgment")

    add_p(
        d,
        "I acknowledge that I have received and reviewed the {{AGENCY_NAME}} Employee Handbook "
        "(the \"Handbook\"). I understand that the Handbook describes important policies and "
        "procedures that apply to my employment, including those addressing equal employment "
        "opportunity, anti-harassment, attendance, time reporting, EVV, drug-free workplace, "
        "social media, and HIPAA.",
    )

    add_p(
        d,
        "I understand that the Handbook is not a contract of employment and does not guarantee "
        "employment for any specific period. My employment is at-will, which means that either I "
        "or the Agency may end the employment relationship at any time, with or without cause and "
        "with or without notice, subject to applicable law. No statement by any Agency employee "
        "can change this at-will relationship except a written agreement signed by an Agency "
        "officer.",
    )

    add_p(
        d,
        "I understand that the Agency may revise the Handbook at any time, and that revisions "
        "supersede prior policies. I am responsible for reading and following the most current "
        "version, available from {{HR_CONTACT_EMAIL}}.",
    )

    add_p(
        d,
        "I agree to comply with the policies in the Handbook and understand that failure to do so "
        "may result in disciplinary action up to and including termination.",
    )

    add_field_line(d, "Handbook version / date received")
    add_signature_block(d, "Employee")

    return save(d, HIRING_DIR / "07-employee-handbook-acknowledgment.docx")


def doc_h08_code_of_conduct() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Code of Conduct & Ethics")

    add_p(
        d,
        "All caregivers, supervisors, contractors, and officers of {{AGENCY_NAME}} are expected "
        "to follow this Code of Conduct. Violations may result in discipline up to and including "
        "termination, and may be reported to licensing authorities, payers, and law enforcement.",
    )

    add_h2(d, "1. Professional Boundaries")
    add_bullets(
        d,
        [
            "Treat clients with dignity and respect at all times.",
            "Do not engage in personal, romantic, or sexual relationships with clients or members of their household.",
            "Do not borrow from, lend to, or co-sign for clients.",
            "Do not bring family members, friends, or pets into a client's home without express Agency authorization.",
            "Do not use a client's car, phone, internet, or supplies for personal purposes.",
        ],
    )

    add_h2(d, "2. Gifts, Tips, and Inheritances")
    add_bullets(
        d,
        [
            "Do not solicit or accept tips, gifts, or money from clients or families, except for nominal items (e.g., holiday cards, baked goods) reportable to your supervisor.",
            "Do not accept inheritances, life insurance designations, or named beneficiary status from clients.",
            "Do not become a power of attorney, guardian, executor, or trustee for a client.",
        ],
    )

    add_h2(d, "3. Honesty in Documentation and Billing")
    add_p(
        d,
        "Documenting visits, time, mileage, and tasks accurately is a legal obligation. "
        "Falsifying EVV records, billing for time not worked, or claiming services not provided "
        "is fraud and may violate the federal False Claims Act (31 U.S.C. § 3729 et seq.) "
        "and state law.",
    )

    add_h2(d, "4. Mandated Reporting")
    add_p(
        d,
        "Caregivers are mandated reporters under {{AGENCY_STATE}} law. Suspected abuse, neglect, "
        "exploitation, or self-neglect involving older adults must be reported to Adult "
        "Protective Services at {{APS_PHONE}}. Suspected child abuse must be reported under "
        "applicable state law (in PA, ChildLine 1-800-932-0313). Make required reports promptly "
        "and notify the Compliance Officer at {{COMPLIANCE_OFFICER_EMAIL}}.",
    )

    add_h2(d, "5. Conflicts of Interest")
    add_bullets(
        d,
        [
            "Disclose any outside employment, ownership interest, or family relationship that could conflict with Agency duties.",
            "Do not accept compensation from a client outside the Agency's billing process.",
            "Do not refer clients to outside services in which you have a financial interest.",
        ],
    )

    add_h2(d, "6. Substance Use, Weapons, Smoking")
    add_bullets(
        d,
        [
            "Do not work under the influence of alcohol, illegal drugs, marijuana, or any substance that impairs your ability to perform safely.",
            "Do not bring weapons of any kind to a client's home or to the Agency office.",
            "Do not smoke, vape, or use tobacco products in a client's home or vehicle.",
        ],
    )

    add_h2(d, "7. Discrimination, Harassment, Retaliation")
    add_p(
        d,
        "The Agency prohibits discrimination and harassment based on any protected status and "
        "prohibits retaliation against anyone who reports a concern in good faith.",
    )

    add_h2(d, "8. Reporting Violations")
    add_p(
        d,
        "Report suspected violations to your supervisor or directly to the Compliance Officer "
        "({{COMPLIANCE_OFFICER_NAME}}, {{COMPLIANCE_OFFICER_EMAIL}}). Reports may be made "
        "anonymously where state law permits.",
    )

    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "08-code-of-conduct-and-ethics.docx")


def doc_h09_hipaa_training() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "HIPAA Training Acknowledgment")

    add_p(
        d,
        "I acknowledge that I completed the {{AGENCY_NAME}} HIPAA Privacy and Security training. "
        "I understand the topics covered, including:",
    )
    add_bullets(
        d,
        [
            "What constitutes Protected Health Information (PHI) and what makes information \"identifiable\" under 45 CFR § 164.514.",
            "Permitted uses and disclosures of PHI for treatment, payment, and operations under 45 CFR § 164.506.",
            "When written authorization is required (45 CFR § 164.508).",
            "Minimum-necessary standard.",
            "Patient rights (access, amendment, accounting, restriction, confidential communication).",
            "Breach notification under 45 CFR Part 164 Subpart D.",
            "Mobile device, EVV device, and home-network security.",
            "Social media restrictions and prohibited disclosures.",
            "Sanctions for violations and the criminal liability under 42 U.S.C. § 1320d-6.",
        ],
    )

    add_p(d, "I will complete annual refresher training as required by Agency policy.")

    add_h2(d, "Training Records")
    add_field_line(d, "Initial training completed (date)")
    add_field_line(d, "Trainer name / title")
    add_field_line(d, "Annual refresher 1 (date)")
    add_field_line(d, "Annual refresher 2 (date)")
    add_field_line(d, "Annual refresher 3 (date)")

    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "09-hipaa-training-acknowledgment.docx")


def doc_h10_mandatory_reporter() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Mandatory Reporter Acknowledgment")

    add_p(
        d,
        "I acknowledge that, as a caregiver employed by {{AGENCY_NAME}}, I am a mandated reporter "
        "of suspected abuse, neglect, and exploitation under {{AGENCY_STATE}} law.",
    )

    add_h2(d, "1. What I Must Report")
    add_bullets(
        d,
        [
            "Physical, sexual, emotional, or verbal abuse of a client.",
            "Neglect (including self-neglect) that places a client at risk of serious harm.",
            "Financial exploitation, theft, or misappropriation of a client's property.",
            "Abandonment.",
            "Suspected child abuse observed in the home.",
            "Conditions in the home that present an immediate risk to the client (e.g., unsafe temperature, lack of food, infestation).",
        ],
    )

    add_h2(d, "2. Where I Report")
    add_bullets(
        d,
        [
            "Adult Protective Services: {{APS_PHONE}}, {{APS_WEBSITE}}",
            "If a child is in immediate danger or you suspect child abuse, call your state hotline (in PA: ChildLine 1-800-932-0313).",
            "If there is an immediate risk of serious physical harm: call 911 first.",
            "After contacting authorities, notify the Agency Compliance Officer at {{COMPLIANCE_OFFICER_EMAIL}}.",
        ],
    )

    add_h2(d, "3. Timing")
    add_p(
        d,
        "Reports must be made promptly, which generally means as soon as I have reasonable cause "
        "to suspect abuse, neglect, or exploitation. Some states require an oral report within "
        "specific hours followed by a written report.",
    )

    add_h2(d, "4. Confidentiality and Non-Retaliation")
    add_p(
        d,
        "Reports made in good faith are confidential under state law and are protected from "
        "retaliation by employers, clients, and family members. The Agency will not retaliate "
        "against any caregiver who makes a good-faith report.",
    )

    add_h2(d, "5. Failure to Report")
    add_p(
        d,
        "Failure to report when required may result in discipline, civil liability, and criminal "
        "penalties under state law.",
    )

    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "10-mandatory-reporter-acknowledgment.docx")


def doc_h11_caregiver_jd() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Caregiver Job Description")

    add_h2(d, "Position Summary")
    add_p(
        d,
        "The Caregiver provides non-medical and personal-care services to clients in their homes "
        "under the direction of the supervising nurse and according to each client's plan of "
        "care. The role supports clients' independence, dignity, safety, and well-being while "
        "complying with HIPAA, state regulations, and Agency policies.",
    )

    add_h2(d, "Reports To")
    add_p(d, "Director of Nursing or designated supervisor.")

    add_h2(d, "Essential Duties and Responsibilities")
    add_bullets(
        d,
        [
            "Provide personal care: bathing, dressing, grooming, toileting, and incontinence care.",
            "Assist with mobility and transfers, including the use of assistive devices.",
            "Prepare meals consistent with the client's dietary needs and assist with feeding when authorized.",
            "Perform light housekeeping in the client's living areas: laundry, dishes, surface cleaning, trash.",
            "Provide medication reminders consistent with state scope-of-practice rules.",
            "Document each visit accurately, including time in/out, tasks completed, observations, and changes in condition.",
            "Use the Agency's EVV system at every visit.",
            "Communicate promptly with the supervising nurse about changes in the client's condition or environment.",
            "Comply with HIPAA, infection control standards, and the Agency's Code of Conduct.",
            "Report suspected abuse, neglect, or exploitation as a mandated reporter.",
            "Complete annual training and skills competency requirements.",
        ],
    )

    add_h2(d, "Qualifications")
    add_bullets(
        d,
        [
            "High school diploma, GED, or equivalent.",
            "Successful completion of any Agency or state-required caregiver training (e.g., {{AGENCY_STATE}}-required orientation hours).",
            "Current CPR/BLS certification.",
            "Negative TB screen as required by state law.",
            "Successful state and federal criminal background check.",
            "Negative pre-employment drug screen.",
            "Verification of identity and employment authorization (Form I-9).",
            "Reliable transportation, valid driver's license, and personal auto insurance meeting Agency limits, if driving for work.",
            "Ability to read, write, and communicate effectively in English (or in another language consistent with client needs).",
        ],
    )

    add_h2(d, "Physical Requirements (ADA-Compliant)")
    add_p(
        d,
        "The physical requirements below describe what is necessary to perform the essential "
        "functions of the position. The Agency will provide reasonable accommodations under the "
        "Americans with Disabilities Act and {{AGENCY_STATE}} law except where doing so would "
        "create undue hardship.",
    )
    add_bullets(
        d,
        [
            "Lift, push, or pull up to 50 pounds occasionally with mechanical assistance or a partner.",
            "Stand, walk, bend, kneel, and reach for extended periods.",
            "Hear and respond to safety alarms, telephones, and verbal communication.",
            "See clearly enough to read medication labels, plan-of-care documents, and assistive-device instructions.",
            "Manual dexterity to operate mobile devices, kitchen tools, and medical equipment.",
            "Tolerate exposure to typical household environmental conditions, pets, cleaning chemicals, and bodily fluids using standard precautions.",
        ],
    )

    add_h2(d, "Working Conditions")
    add_bullets(
        d,
        [
            "Work is performed in clients' homes and may involve travel between homes.",
            "Schedules vary, including evenings, weekends, and holidays.",
            "On-call expectations may apply for some positions.",
        ],
    )

    add_h2(d, "Acknowledgement")
    add_p(
        d,
        "I have reviewed this job description, understand the essential functions and physical "
        "requirements of the position, and confirm that I can perform the essential functions "
        "with or without reasonable accommodation.",
    )
    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "11-caregiver-job-description.docx")


def doc_h12_i9_w4_cover() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "New Hire Federal Forms Cover Sheet")

    add_p(
        d,
        "This cover sheet accompanies the federal forms every new employee must complete. The "
        "actual federal forms (Form I-9 and Form W-4) are not reproduced here; the most current "
        "editions must be downloaded from the issuing agency at the time of hire.",
        italic=True,
    )

    add_h2(d, "Form I-9 (USCIS)")
    add_p(d, "Source of current edition: https://www.uscis.gov/i-9 (always use the latest version).")
    add_bullets(
        d,
        [
            "Section 1 must be completed and signed by the employee on or before the first day of work.",
            "Section 2 must be completed by the employer within three (3) business days of the employee's first day of work.",
            "Acceptable documents are listed on the Form I-9 itself. Examples:",
            "List A (proves identity AND work authorization): U.S. passport; Permanent Resident Card; Employment Authorization Document.",
            "List B (proves identity only): driver's license; state ID card; school ID with photo.",
            "List C (proves work authorization only): Social Security card; certified U.S. birth certificate; Form I-94.",
            "An employee must present ONE document from List A, OR one from List B and one from List C.",
            "Photocopy and retain only as required by your I-9 retention policy.",
        ],
    )

    add_h2(d, "Form W-4 (IRS)")
    add_p(d, "Source of current edition: https://www.irs.gov/forms-pubs/about-form-w-4 (always use the latest version).")
    add_bullets(
        d,
        [
            "Required for federal income tax withholding.",
            "Employee completes and signs.",
            "If no Form W-4 is submitted, withhold as if the employee is single with no other adjustments.",
        ],
    )

    add_h2(d, "State Withholding")
    add_p(
        d,
        "Pennsylvania uses Form REV-419 (Employee's Nonwithholding Application Certificate) where "
        "applicable. Ohio uses Form IT-4. Local taxes may also apply (PA Local Earned Income Tax, "
        "various Ohio municipal taxes). Confirm at hire with HR ({{HR_CONTACT_EMAIL}}).",
    )

    add_h2(d, "New Hire Reporting")
    add_p(
        d,
        "The Agency reports each new hire to the {{AGENCY_STATE}} New Hire Reporting Center "
        "within twenty (20) days of hire as required by federal law (42 U.S.C. § 653a).",
    )

    add_h2(d, "Onboarding Checklist")
    add_bullets(
        d,
        [
            "[ ] I-9 Section 1 (employee) completed by first day of work.",
            "[ ] I-9 Section 2 (employer) completed within three business days.",
            "[ ] Form W-4 received.",
            "[ ] State withholding form received (PA REV-419 / OH IT-4 / other).",
            "[ ] Direct Deposit Authorization received with voided check or bank letter.",
            "[ ] Background-check authorization signed; results received and reviewed.",
            "[ ] Drug-screen results received.",
            "[ ] License/certification verifications complete.",
            "[ ] Confidentiality / NDA signed.",
            "[ ] Employee Handbook acknowledgment signed.",
            "[ ] HIPAA training acknowledgment signed.",
            "[ ] Mandatory reporter acknowledgment signed.",
            "[ ] Code of Conduct signed.",
            "[ ] Caregiver job description signed.",
            "[ ] Driver / vehicle verification (if applicable) complete.",
            "[ ] New-hire reported to state.",
        ],
    )

    add_field_line(d, "Onboarding completed by (HR)")
    add_field_line(d, "Date")
    return save(d, HIRING_DIR / "12-i9-w4-new-hire-cover-sheet.docx")


def doc_h13_driver_vehicle() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Driver and Vehicle Verification")

    add_p(
        d,
        "Caregivers who drive personal vehicles for work, transport clients, or run client errands "
        "must complete this verification before any driving on behalf of {{AGENCY_NAME}}.",
    )

    add_h2(d, "1. Driver's License")
    add_field_line(d, "Full legal name as it appears on license")
    add_field_line(d, "License number")
    add_field_line(d, "State of issuance")
    add_field_line(d, "License expiration")
    add_field_line(d, "Class / endorsements / restrictions")

    add_h2(d, "2. Motor Vehicle Record (MVR) Authorization")
    add_p(
        d,
        "I authorize {{AGENCY_NAME}} or its designated vendor to obtain my motor vehicle record "
        "from {{AGENCY_STATE}} and any other state where I have held a license, before hire and "
        "annually thereafter. I understand that an unsatisfactory MVR may disqualify me from "
        "driving on behalf of the Agency.",
    )

    add_h2(d, "3. Vehicle and Insurance")
    add_field_line(d, "Vehicle make / model / year")
    add_field_line(d, "License plate / state")
    add_field_line(d, "Insurance carrier")
    add_field_line(d, "Policy number")
    add_field_line(d, "Effective dates")
    add_p(
        d,
        "Minimum required liability limits: {{INSURANCE_AUTO_LIMIT}}. (Bodily injury per person / "
        "bodily injury per occurrence / property damage.) The Agency may require higher limits "
        "for caregivers who transport clients.",
    )
    add_p(
        d,
        "Attach a copy of the current declarations page or insurance ID card. The caregiver is "
        "responsible for notifying the Agency in writing within five (5) business days of any "
        "lapse or change in coverage.",
    )

    add_h2(d, "4. Use of Personal Vehicle for Work")
    add_p(
        d,
        "Mileage between client visits is reimbursed at {{MILEAGE_RATE}} per mile in accordance "
        "with the Agency's mileage policy. Commuting between home and the first/last visit of the "
        "day is generally not reimbursed unless the Agency specifies otherwise. Personal use of a "
        "vehicle while on duty is prohibited.",
    )

    add_h2(d, "5. Safe Driving and Mobile Devices")
    add_bullets(
        d,
        [
            "Comply with all traffic laws.",
            "Do not text, email, browse, or hand-hold a phone while operating a vehicle.",
            "Pull over before responding to client or Agency calls.",
            "Do not transport unauthorized passengers in a vehicle being used for work.",
            "Wear seatbelts and require all passengers to do so.",
        ],
    )

    add_h2(d, "6. Reportable Events")
    add_p(
        d,
        "I will notify the Agency within twenty-four (24) hours of any motor vehicle accident, "
        "moving violation, license suspension, or DUI/DWI arrest or conviction.",
    )

    add_signature_block(d, "Employee")
    return save(d, HIRING_DIR / "13-driver-and-vehicle-verification.docx")


# ===========================================================================
# STATE ADDENDA (2)
# ===========================================================================


def doc_addendum_pa() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Pennsylvania State Addendum")
    add_subtitle(d, "Attach to the master service agreement and hiring packet")

    add_p(
        d,
        "This addendum supplements the {{AGENCY_NAME}} client and employment documents with "
        "Pennsylvania-specific requirements. Where this addendum conflicts with a general "
        "provision, this addendum controls for services delivered or work performed in "
        "Pennsylvania.",
    )

    add_h2(d, "1. Home Care Agency Licensure")
    add_p(
        d,
        "The Agency operates as a home care agency and/or home care registry licensed under 28 "
        "Pa. Code Chapter 611, administered by the Pennsylvania Department of Health, Division "
        "of Home Health (telephone {{PA_DEPT_HEALTH_PHONE}}). Agency license number: "
        "{{PA_HOMECARE_LICENSE}}.",
    )

    add_h2(d, "2. Older Adults Protective Services Act (OAPSA)")
    add_p(
        d,
        "The Agency complies with the Older Adults Protective Services Act, 35 P.S. § "
        "10225.101 et seq., including criminal history record checks for direct-care employees "
        "prior to hire, FBI fingerprint checks where required, and the prohibitions and "
        "procedures in 35 P.S. § 10225.503. Suspected abuse, neglect, exploitation, or "
        "abandonment of an older adult will be reported to the Pennsylvania Department of Aging "
        "or the local Area Agency on Aging at {{APS_PHONE}}.",
    )

    add_h2(d, "3. Child Protective Services Law")
    add_p(
        d,
        "The Agency complies with the Pennsylvania Child Protective Services Law (Acts 13 and "
        "151 of 2014), including required clearances (PA State Police criminal history, PA Child "
        "Abuse History, FBI fingerprint) and mandated reporting. Suspected child abuse will be "
        "reported promptly to ChildLine at 1-800-932-0313.",
    )

    add_h2(d, "4. Wage Payment")
    add_p(
        d,
        "The Agency complies with the Pennsylvania Wage Payment and Collection Law, 43 P.S. "
        "§ 260.1 et seq., including timely payment, lawful deductions, and final-pay "
        "requirements.",
    )

    add_h2(d, "5. PA Human Relations Act")
    add_p(
        d,
        "The Agency complies with the Pennsylvania Human Relations Act, 43 P.S. § 951 et "
        "seq. Discrimination and harassment in employment or services on a protected basis is "
        "prohibited.",
    )

    add_h2(d, "6. Advance Directives")
    add_p(
        d,
        "Advance directives, including living wills and health care powers of attorney, are "
        "recognized under Act 169 of 2002 (20 Pa. C.S. Ch. 54). The Agency will follow advance "
        "directives that comply with Pennsylvania law.",
    )

    add_h2(d, "7. Workers' Compensation")
    add_p(
        d,
        "Employees are covered by workers' compensation through {{WORKERS_COMP_CARRIER}}, "
        "consistent with the Pennsylvania Workers' Compensation Act, 77 P.S. § 1 et seq.",
    )

    add_h2(d, "8. Complaints")
    add_p(
        d,
        "Concerns about Agency services may be made directly to the Pennsylvania Department of "
        "Health Home Health Hotline ({{PA_DEPT_HEALTH_PHONE}}) or to the Pennsylvania Department "
        "of Aging at {{APS_PHONE}}.",
    )

    add_p(
        d,
        "By signing the master service agreement and/or employment documents, the parties "
        "acknowledge they have read and accept this Pennsylvania addendum.",
    )

    add_two_party_signatures(d, "Client / Employee", "Agency representative")
    return save(d, ADDENDA_DIR / "pennsylvania-addendum.docx")


def doc_addendum_oh() -> Path:
    d = new_doc()
    add_letterhead(d)
    add_title(d, "Ohio State Addendum")
    add_subtitle(d, "Attach to the master service agreement and hiring packet")

    add_p(
        d,
        "This addendum supplements the {{AGENCY_NAME}} client and employment documents with "
        "Ohio-specific requirements. Where this addendum conflicts with a general provision, "
        "this addendum controls for services delivered or work performed in Ohio.",
    )

    add_h2(d, "1. Home Health and Home Care Regulation")
    add_p(
        d,
        "The Agency operates as a home health agency or home care provider as regulated under "
        "Ohio Revised Code Chapter 3740, with related rules at OAC Chapter 3701-13. The Ohio "
        "Department of Health (telephone {{OH_DEPT_HEALTH_PHONE}}) is the principal regulator. "
        "Agency provider number: {{OH_PROVIDER_NUMBER}}.",
    )
    add_p(
        d,
        "If the Agency participates in the PASSPORT waiver, Choices, Assisted Living, or other "
        "Aging Network programs, ORC Chapter 173 (Department of Aging) and the related "
        "administrative rules also apply.",
    )

    add_h2(d, "2. Criminal Records Checks")
    add_p(
        d,
        "The Agency complies with ORC § 109.572 and OAC Chapter 3701-13 requirements for "
        "criminal records checks, including BCI&I and FBI fingerprint checks for direct-care "
        "employees. Disqualifying offenses listed in OAC 3701-13-04 are applied as required.",
    )

    add_h2(d, "3. Adult Protective Services")
    add_p(
        d,
        "Suspected abuse, neglect, or exploitation of an adult age sixty (60) or older will be "
        "reported under ORC § 5101.61 to the local county Department of Job and Family "
        "Services or to the statewide Adult Protective Services hotline at 1-855-OHIO-APS "
        "(1-855-644-6277). Caregivers are mandated reporters.",
    )

    add_h2(d, "4. Child Protective Services")
    add_p(
        d,
        "Suspected child abuse will be reported under ORC § 2151.421 to the local Public "
        "Children Services Agency or law enforcement.",
    )

    add_h2(d, "5. Ohio Civil Rights Act")
    add_p(
        d,
        "The Agency complies with the Ohio Civil Rights Act, ORC Chapter 4112. Discrimination "
        "and harassment in employment or services on a protected basis is prohibited.",
    )

    add_h2(d, "6. Wage Payment")
    add_p(
        d,
        "The Agency complies with Ohio's wage payment requirements (ORC Ch. 4111 and 4113), "
        "including timely payment of wages and lawful deductions.",
    )

    add_h2(d, "7. Workers' Compensation")
    add_p(
        d,
        "Employees are covered by Ohio Bureau of Workers' Compensation coverage through "
        "{{WORKERS_COMP_CARRIER}}, consistent with ORC Chapter 4123.",
    )

    add_h2(d, "8. Advance Directives")
    add_p(
        d,
        "Advance directives are recognized under ORC Chapters 1337 and 2133 (durable power of "
        "attorney for health care; living will). The Agency will follow advance directives that "
        "comply with Ohio law.",
    )

    add_h2(d, "9. Complaints")
    add_p(
        d,
        "Concerns about Agency services may be made directly to the Ohio Department of Health at "
        "{{OH_DEPT_HEALTH_PHONE}} or to Adult Protective Services at 1-855-OHIO-APS.",
    )

    add_p(
        d,
        "By signing the master service agreement and/or employment documents, the parties "
        "acknowledge they have read and accept this Ohio addendum.",
    )

    add_two_party_signatures(d, "Client / Employee", "Agency representative")
    return save(d, ADDENDA_DIR / "ohio-addendum.docx")


# ===========================================================================
# Driver
# ===========================================================================


GENERATORS = [
    doc_01_homecare_services_agreement,
    doc_02_hipaa_notice_of_privacy_practices,
    doc_03_hipaa_authorization,
    doc_04_client_bill_of_rights,
    doc_05_consent_to_care,
    doc_06_financial_responsibility,
    doc_07_advance_directives,
    doc_08_photo_media_release,
    doc_09_emergency_contact,
    doc_10_grievance,
    doc_h01_employment_application,
    doc_h02_offer_letter,
    doc_h03_confidentiality_nda,
    doc_h04_background_check_fcra,
    doc_h05_drug_alcohol,
    doc_h06_direct_deposit,
    doc_h07_handbook_ack,
    doc_h08_code_of_conduct,
    doc_h09_hipaa_training,
    doc_h10_mandatory_reporter,
    doc_h11_caregiver_jd,
    doc_h12_i9_w4_cover,
    doc_h13_driver_vehicle,
    doc_addendum_pa,
    doc_addendum_oh,
]


def main() -> int:
    print(f"Generating {len(GENERATORS)} master .docx files...")
    print(f"Output root: {ROOT}")
    print()
    for fn in GENERATORS:
        path = fn()
        rel = path.relative_to(ROOT)
        print(f"  wrote {rel}")
    print()
    print(f"Done. {len(GENERATORS)} files generated.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
