"""
Customize all legal templates for a specific homecare agency.

This script reads `agency-info.json`, performs find-and-replace on every
{{PLACEHOLDER}} across all .docx templates, and exports filled .docx and
.pdf copies into a per-agency subfolder so the originals stay untouched.

Usage:
    # 1. Edit agency-info.json (or have Gemini fill it out)
    # 2. Run:
    python3 scripts/customize.py
    # or with a custom config path:
    python3 scripts/customize.py path/to/my-agency.json

Output:
    customized/<agency-slug>/docx/*.docx   <-- ready to email or print
    customized/<agency-slug>/pdf/*.pdf     <-- ready to send to clients/staff
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CONFIG = ROOT / "agency-info.json"


def slugify(text: str) -> str:
    text = re.sub(r"[^a-zA-Z0-9\s-]", "", text).strip().lower()
    return re.sub(r"[\s_]+", "-", text) or "agency"


def replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving runs/formatting where possible.

    python-docx splits text across runs so a {{PLACEHOLDER}} can span runs.
    Strategy: if a placeholder appears in the joined text, rewrite the whole
    paragraph text on the first run and clear the rest. We accept loss of
    inline run-level formatting inside replaced placeholders — this is fine
    because the templates don't use mid-placeholder formatting.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    for key, value in mapping.items():
        new_text = new_text.replace("{{" + key + "}}", value)

    if new_text == full_text:
        return

    # Write the whole replaced text into the first run, clear the others.
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_in_doc(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, mapping)


def collect_templates() -> list[Path]:
    paths: list[Path] = []
    for folder in ("client-onboarding", "hiring", "state-addenda"):
        paths.extend(sorted((ROOT / folder / "docx").glob("*.docx")))
    return paths


def docx_to_pdf(docx_path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx to .pdf via libreoffice headless. Returns the PDF path or None."""
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        print(f"  WARN: PDF conversion failed for {docx_path.name}: {e}")
        return None
    return out_dir / (docx_path.stem + ".pdf")


def main() -> None:
    cfg_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_CONFIG
    if not cfg_path.exists():
        print(f"ERROR: Config file not found: {cfg_path}")
        print("Edit agency-info.json with your agency's information, then re-run.")
        sys.exit(1)

    cfg = json.loads(cfg_path.read_text())

    # Stringify everything for safe replacement
    mapping = {k: ("" if v is None else str(v)) for k, v in cfg.items()}

    agency_name = mapping.get("AGENCY_NAME", "agency").strip() or "agency"
    slug = slugify(agency_name)

    out_root = ROOT / "customized" / slug
    out_docx = out_root / "docx"
    out_pdf = out_root / "pdf"
    if out_root.exists():
        shutil.rmtree(out_root)
    out_docx.mkdir(parents=True, exist_ok=True)
    out_pdf.mkdir(parents=True, exist_ok=True)

    templates = collect_templates()
    print(f"Customizing {len(templates)} templates for: {agency_name}")
    print(f"Output: {out_root.relative_to(ROOT)}/")
    print()

    for tpl in templates:
        category = tpl.parent.parent.name  # client-onboarding | hiring | state-addenda
        category_docx = out_docx / category
        category_pdf = out_pdf / category
        category_docx.mkdir(parents=True, exist_ok=True)
        category_pdf.mkdir(parents=True, exist_ok=True)

        doc = Document(tpl)
        replace_in_doc(doc, mapping)
        out_file = category_docx / tpl.name
        doc.save(out_file)

        pdf = docx_to_pdf(out_file, category_pdf)
        marker = "PDF" if pdf else "docx-only"
        print(f"  [{marker}] {category}/{tpl.stem}")

    print()
    print("Done.")
    print(f"Filled DOCX: {out_docx.relative_to(ROOT)}/")
    print(f"Filled PDF:  {out_pdf.relative_to(ROOT)}/")


if __name__ == "__main__":
    main()
