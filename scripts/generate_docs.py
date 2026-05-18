"""
Generate all legal onboarding & hiring .docx templates for homecare agencies.

Templates use {{PLACEHOLDER}} merge fields so an agency (or Gemini) can
do a simple find-and-replace to customize.

Run:
    python3 scripts/generate_docs.py

Outputs:
    client-onboarding/docx/*.docx
    hiring/docx/*.docx
    state-addenda/docx/*.docx
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CLIENT_DIR = ROOT / "client-onboarding" / "docx"
HIRING_DIR = ROOT / "hiring" / "docx"
STATE_DIR = ROOT / "state-addenda" / "docx"


# -----------------------------
# Document-building primitives
# -----------------------------


def new_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    return doc


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_signature_block(doc: Document, role_label: str = "Client / Authorized Representative") -> None:
    doc.add_paragraph()
    add_h2(doc, "Signatures")
    doc.add_paragraph(f"{role_label} Signature: ___________________________________   Date: ____________")
    doc.add_paragraph(f"Printed Name: ___________________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Agency Representative Signature: __________________________   Date: ____________")
    doc.add_paragraph("Printed Name & Title: ___________________________________")


def agency_header(doc: Document) -> None:
    """Standard editable header that appears on every document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{AGENCY_NAME}}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("{{AGENCY_ADDRESS}} | {{AGENCY_CITY}}, {{AGENCY_STATE}} {{AGENCY_ZIP}}").font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Phone: {{AGENCY_PHONE}} | Email: {{AGENCY_EMAIL}} | License #: {{AGENCY_LICENSE_NUMBER}}").font.size = Pt(10)

    doc.add_paragraph()


def save(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"  wrote {path.relative_to(ROOT)}")


# -----------------------------
# CLIENT ONBOARDING DOCUMENTS
# -----------------------------


def doc_service_agreement() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "HOMECARE SERVICES AGREEMENT")
    add_subtitle(doc, "This agreement governs the services provided by {{AGENCY_NAME}} to the Client identified below.")

    add_h2(doc, "1. Parties")
    add_p(doc, "This Homecare Services Agreement (\"Agreement\") is entered into as of {{EFFECTIVE_DATE}} between {{AGENCY_NAME}} (\"Agency\"), a homecare agency licensed in the State of {{AGENCY_STATE}}, and {{CLIENT_NAME}} (\"Client\"), residing at {{CLIENT_ADDRESS}}.")

    add_h2(doc, "2. Services Provided")
    add_p(doc, "Agency will provide non-medical and/or skilled homecare services as described in the Client's individualized Plan of Care, which is incorporated into this Agreement by reference. Services may include, but are not limited to:")
    add_bullets(doc, [
        "Personal care assistance (bathing, grooming, dressing, toileting)",
        "Mobility and transfer assistance",
        "Meal planning and preparation",
        "Light housekeeping and laundry",
        "Medication reminders (per state scope of practice)",
        "Companionship and supervision",
        "Transportation and errand assistance (where authorized)",
        "Skilled nursing or therapy services (where licensed and ordered)",
    ])

    add_h2(doc, "3. Service Schedule and Hours")
    add_p(doc, "Services will be provided on the schedule documented in the Plan of Care. The agreed hourly rate is ${{HOURLY_RATE}} per hour, with a minimum visit length of {{MIN_VISIT_HOURS}} hours. Holiday, overnight, and live-in rates are listed in the Financial Responsibility Agreement.")

    add_h2(doc, "4. Electronic Visit Verification (EVV)")
    add_p(doc, "Agency uses an Electronic Visit Verification system to record the date, time, location, and tasks of each caregiver visit, in accordance with the 21st Century Cures Act and applicable state Medicaid requirements. Client consents to EVV check-in/check-out from the Client's residence and authorizes geolocation capture solely for visit verification.")

    add_h2(doc, "5. Caregiver Assignment")
    add_p(doc, "Agency will make reasonable efforts to assign qualified caregivers consistent with the Client's needs and preferences. Agency reserves the right to substitute caregivers as needed. All caregivers are employees or contractors of Agency, are screened in accordance with state law, and are bound by confidentiality obligations.")

    add_h2(doc, "6. Client Responsibilities")
    add_bullets(doc, [
        "Provide a safe work environment, free from harassment, hazardous conditions, and unrestrained pets that pose a risk.",
        "Notify Agency at least {{CANCEL_NOTICE_HOURS}} hours in advance of any visit cancellation; late cancellations may incur a fee.",
        "Refrain from giving caregivers cash, gifts of significant value, loans, or access to financial accounts, vehicles, or weapons.",
        "Provide accurate health information, including medications, allergies, and emergency contacts.",
        "Treat caregivers and Agency staff with respect, free from discrimination or harassment of any kind.",
    ])

    add_h2(doc, "7. Payment Terms")
    add_p(doc, "Invoices are issued {{BILLING_FREQUENCY}} and are due {{PAYMENT_TERMS_DAYS}} days from the invoice date. Accepted forms of payment: {{PAYMENT_METHODS}}. Past-due balances over {{LATE_THRESHOLD_DAYS}} days may incur a late fee of {{LATE_FEE_PCT}}% per month and may result in suspension of services. Client is responsible for any private-pay balances not covered by insurance, Medicaid, VA, or long-term care benefits.")

    add_h2(doc, "8. Termination")
    add_p(doc, "Either party may terminate this Agreement with {{TERMINATION_NOTICE_DAYS}} days' written notice. Agency may terminate immediately for cause, including but not limited to: unsafe conditions, abuse or harassment of caregivers, non-payment beyond {{NONPAYMENT_THRESHOLD_DAYS}} days, fraud, or violation of this Agreement. Upon termination, Client remains responsible for all services rendered through the termination date.")

    add_h2(doc, "9. Limitation of Liability")
    add_p(doc, "Agency carries general liability and workers' compensation insurance. To the maximum extent permitted by law, Agency's total liability under this Agreement is limited to the fees paid by Client in the 90 days preceding the event giving rise to the claim. Agency is not liable for any indirect, incidental, or consequential damages.")

    add_h2(doc, "10. Governing Law and Disputes")
    add_p(doc, "This Agreement is governed by the laws of the State of {{AGENCY_STATE}}. The parties agree to attempt resolution through good-faith mediation before pursuing litigation. Venue for any litigation shall be the courts of {{AGENCY_COUNTY}} County, {{AGENCY_STATE}}.")

    add_h2(doc, "11. Entire Agreement")
    add_p(doc, "This Agreement, together with the Plan of Care, HIPAA Notice, Financial Responsibility Agreement, Bill of Rights, and any state-specific addendum, constitutes the entire agreement between the parties. Modifications must be in writing and signed by both parties.")

    add_signature_block(doc)
    return doc


def doc_hipaa_notice() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "NOTICE OF PRIVACY PRACTICES (HIPAA)")
    add_subtitle(doc, "Effective Date: {{EFFECTIVE_DATE}}")

    add_p(doc, "THIS NOTICE DESCRIBES HOW MEDICAL INFORMATION ABOUT YOU MAY BE USED AND DISCLOSED AND HOW YOU CAN GET ACCESS TO THIS INFORMATION. PLEASE REVIEW IT CAREFULLY.")

    add_h2(doc, "Our Commitment to Your Privacy")
    add_p(doc, "{{AGENCY_NAME}} is required by law to maintain the privacy of your Protected Health Information (\"PHI\"), provide you with this Notice of our legal duties and privacy practices, and abide by the terms of the Notice currently in effect.")

    add_h2(doc, "How We May Use and Disclose Your PHI Without Your Authorization")
    add_bullets(doc, [
        "Treatment: To provide, coordinate, or manage your care with caregivers, nurses, physicians, and other providers.",
        "Payment: To bill Medicare, Medicaid, private insurance, VA, long-term care insurers, or you directly.",
        "Health Care Operations: For quality improvement, training, audits, accreditation, and business management.",
        "As Required by Law: For public health reporting, abuse/neglect reporting, court orders, law enforcement, and oversight agencies.",
        "To Avert a Serious Threat: To prevent or lessen a serious threat to health or safety.",
    ])

    add_h2(doc, "Uses and Disclosures Requiring Your Written Authorization")
    add_p(doc, "We will obtain your written authorization before using or disclosing your PHI for marketing, sale of PHI, most psychotherapy notes, or any purpose not described in this Notice. You may revoke an authorization in writing at any time, except to the extent we have already acted in reliance on it.")

    add_h2(doc, "Your Rights Regarding Your PHI")
    add_bullets(doc, [
        "Right to inspect and obtain a copy of your records.",
        "Right to request an amendment of your PHI.",
        "Right to an accounting of certain disclosures we have made.",
        "Right to request restrictions on uses and disclosures.",
        "Right to request confidential communications by alternative means or location.",
        "Right to receive a paper copy of this Notice on request.",
        "Right to be notified of any breach of unsecured PHI.",
    ])

    add_h2(doc, "Complaints")
    add_p(doc, "If you believe your privacy rights have been violated, you may file a complaint with our Privacy Officer at {{PRIVACY_OFFICER_NAME}}, {{AGENCY_PHONE}}, {{PRIVACY_OFFICER_EMAIL}}, or with the U.S. Department of Health and Human Services Office for Civil Rights. We will not retaliate against you for filing a complaint.")

    add_h2(doc, "Changes to This Notice")
    add_p(doc, "We reserve the right to change this Notice and to make the new Notice effective for all PHI we maintain. Updated Notices will be posted at our office and on our website.")

    add_h2(doc, "Acknowledgment of Receipt")
    add_p(doc, "I acknowledge that I have received a copy of {{AGENCY_NAME}}'s Notice of Privacy Practices.")
    add_signature_block(doc)
    return doc


def doc_hipaa_authorization() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "AUTHORIZATION FOR USE / DISCLOSURE OF PROTECTED HEALTH INFORMATION")

    add_p(doc, "Client Name: {{CLIENT_NAME}}    Date of Birth: {{CLIENT_DOB}}")
    add_p(doc, "I authorize {{AGENCY_NAME}} to use and/or disclose the Protected Health Information (PHI) described below.")

    add_h2(doc, "Information to Be Used or Disclosed")
    add_bullets(doc, [
        "Plan of care, progress notes, and visit documentation",
        "Medication list and administration records",
        "Physician orders and discharge summaries",
        "Billing and insurance information",
        "Other (specify): _________________________________________________",
    ])

    add_h2(doc, "Person(s) / Entity Authorized to Receive the Information")
    add_p(doc, "Name: ____________________________________  Relationship: _______________")
    add_p(doc, "Address: __________________________________________________________________")
    add_p(doc, "Phone: ____________________________  Email: ____________________________")

    add_h2(doc, "Purpose of the Disclosure")
    add_p(doc, "[ ] Care coordination   [ ] Family communication   [ ] Insurance/Billing   [ ] Legal   [ ] Other: ________________")

    add_h2(doc, "Expiration")
    add_p(doc, "This authorization will expire on {{AUTHORIZATION_EXPIRATION_DATE}} or upon the following event: ___________________________. If no date is specified, this authorization expires one (1) year from the date signed.")

    add_h2(doc, "Right to Revoke")
    add_p(doc, "I understand I may revoke this authorization in writing at any time, except to the extent the Agency has already acted on it. Revocations should be sent to {{PRIVACY_OFFICER_EMAIL}}.")

    add_h2(doc, "Re-disclosure")
    add_p(doc, "I understand that information disclosed under this authorization may be re-disclosed by the recipient and may no longer be protected by HIPAA.")

    add_signature_block(doc)
    return doc


def doc_bill_of_rights() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "CLIENT BILL OF RIGHTS AND RESPONSIBILITIES")

    add_p(doc, "{{AGENCY_NAME}} is committed to protecting the rights and dignity of every client. As a client of this Agency, you have the following rights:")

    add_h2(doc, "Your Rights")
    add_bullets(doc, [
        "To be treated with consideration, respect, and dignity, free from abuse, neglect, exploitation, or discrimination on the basis of race, color, religion, national origin, sex, sexual orientation, gender identity, age, disability, marital status, or source of payment.",
        "To receive information about services, costs, and any limitations in advance.",
        "To participate in the development of your Plan of Care and to refuse any service or treatment.",
        "To privacy and confidentiality of your health information, in accordance with HIPAA.",
        "To voice grievances or recommend changes in services without fear of retaliation.",
        "To be informed of caregiver qualifications and to request a different caregiver.",
        "To be informed of the Agency's policies on advance directives.",
        "To be informed of any financial relationships between the Agency and other providers.",
        "To receive services in your home environment without unauthorized recording or surveillance.",
        "To have your property treated with respect.",
        "To be free from physical or chemical restraints used for discipline or staff convenience.",
    ])

    add_h2(doc, "Your Responsibilities")
    add_bullets(doc, [
        "Provide accurate health information and notify the Agency of any changes.",
        "Follow the agreed-upon Plan of Care or notify Agency if you cannot.",
        "Provide a safe, harassment-free environment for caregivers.",
        "Pay for services as agreed and notify Agency promptly of insurance changes.",
        "Treat caregivers and staff with respect.",
    ])

    add_h2(doc, "Reporting Concerns")
    add_p(doc, "If you have a concern, contact: {{COMPLIANCE_OFFICER_NAME}}, {{AGENCY_PHONE}}, {{COMPLIANCE_OFFICER_EMAIL}}.")
    add_p(doc, "You may also report concerns to your state regulator: {{STATE_REGULATOR_NAME}} — {{STATE_REGULATOR_PHONE}}.")
    add_p(doc, "Adult Protective Services: {{APS_PHONE}}.")

    add_signature_block(doc)
    return doc


def doc_consent_to_care() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "CONSENT TO CARE")

    add_p(doc, "Client Name: {{CLIENT_NAME}}    Date of Birth: {{CLIENT_DOB}}")
    add_p(doc, "I voluntarily consent to receive homecare services from {{AGENCY_NAME}} as described in my Plan of Care. I understand that:")
    add_bullets(doc, [
        "Services will be provided by qualified caregivers employed or contracted by the Agency.",
        "I have the right to refuse any service at any time.",
        "I have the right to be informed of and to participate in decisions regarding my care.",
        "Services may include, but are not limited to, personal care, mobility assistance, meal preparation, medication reminders, light housekeeping, companionship, and (where ordered) skilled nursing.",
        "If I am unable to provide consent, my legally authorized representative may consent on my behalf.",
        "I authorize the Agency to contact my physician(s) and other providers as needed to coordinate care.",
        "I understand that homecare carries inherent risks and that no specific outcome is guaranteed.",
    ])

    add_h2(doc, "Emergency Care")
    add_p(doc, "In the event of a medical emergency, I authorize the Agency's caregivers to call 911 and to transport or arrange transport to the nearest hospital. The Agency will notify my emergency contact as soon as reasonably possible.")

    add_signature_block(doc)
    return doc


def doc_financial_responsibility() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "FINANCIAL RESPONSIBILITY AND PAYMENT AUTHORIZATION")

    add_h2(doc, "Rate Schedule")
    add_p(doc, "Standard hourly rate: ${{HOURLY_RATE}}/hr (minimum {{MIN_VISIT_HOURS}}-hour visit)")
    add_p(doc, "Weekend rate: ${{WEEKEND_RATE}}/hr")
    add_p(doc, "Holiday rate: ${{HOLIDAY_RATE}}/hr")
    add_p(doc, "Overnight (per-shift): ${{OVERNIGHT_RATE}}")
    add_p(doc, "Live-in (per-day): ${{LIVE_IN_RATE}}")
    add_p(doc, "Mileage reimbursement: ${{MILEAGE_RATE}}/mile when authorized")

    add_h2(doc, "Payer Information")
    add_p(doc, "Primary Payer: [ ] Private Pay   [ ] Medicare   [ ] Medicaid/Waiver   [ ] LTC Insurance   [ ] VA   [ ] Other: __________")
    add_p(doc, "Insurance Carrier / Plan: ___________________________  Policy #: ____________________  Group #: ____________")

    add_h2(doc, "Billing Terms")
    add_bullets(doc, [
        "Invoices are issued {{BILLING_FREQUENCY}} for services rendered.",
        "Payment is due within {{PAYMENT_TERMS_DAYS}} days of the invoice date.",
        "Accepted methods: {{PAYMENT_METHODS}}.",
        "Returned-payment fee: ${{RETURNED_PAYMENT_FEE}} per occurrence.",
        "Late fee: {{LATE_FEE_PCT}}% per month on balances over {{LATE_THRESHOLD_DAYS}} days past due.",
        "Cancellations with less than {{CANCEL_NOTICE_HOURS}} hours' notice: full visit fee may be charged.",
    ])

    add_h2(doc, "Payment Authorization (Optional)")
    add_p(doc, "[ ] I authorize {{AGENCY_NAME}} to charge the credit card or debit my bank account on file for invoices when due.")
    add_p(doc, "Card / Account ending in: _______________   Expiration: __________")

    add_h2(doc, "Insurance Assignment of Benefits")
    add_p(doc, "I assign all applicable insurance benefits directly to {{AGENCY_NAME}}. I authorize the Agency to release information necessary to process my claims and to appeal any adverse determinations on my behalf.")

    add_signature_block(doc, role_label="Client / Financially Responsible Party")
    return doc


def doc_advance_directive_ack() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "ADVANCE DIRECTIVES ACKNOWLEDGMENT")

    add_p(doc, "{{AGENCY_NAME}} respects your right to make decisions about your medical care, including the right to accept or refuse treatment and to formulate advance directives.")

    add_h2(doc, "What Are Advance Directives?")
    add_p(doc, "Advance directives are written instructions, recognized under {{AGENCY_STATE}} law, regarding your wishes for medical care if you become unable to speak for yourself. Common advance directives include a Living Will, Durable Power of Attorney for Health Care, and out-of-hospital DNR orders.")

    add_h2(doc, "Your Choices")
    add_p(doc, "[ ] I have an advance directive. A copy is attached / will be provided.")
    add_p(doc, "[ ] I do not have an advance directive and would like more information.")
    add_p(doc, "[ ] I do not have an advance directive and decline information at this time.")

    add_h2(doc, "Agency Policy")
    add_p(doc, "Caregivers are not authorized to witness or sign as a witness to advance directives. The Agency will follow valid advance directives to the extent permitted by law and Agency scope of practice. In the event of a medical emergency without a valid out-of-hospital DNR, caregivers will call 911 and provide CPR if trained.")

    add_signature_block(doc)
    return doc


def doc_photo_media_release() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "PHOTO AND MEDIA RELEASE (OPTIONAL)")

    add_p(doc, "I, {{CLIENT_NAME}}, give {{AGENCY_NAME}} permission to use my likeness, photographs, video, and/or testimonial in the following ways (check all that apply):")
    add_bullets(doc, [
        "[ ] Internal training and quality-improvement materials",
        "[ ] Agency website and social media",
        "[ ] Marketing brochures, advertisements, and press materials",
        "[ ] News and media features",
    ])
    add_p(doc, "I understand that I will not receive compensation. I may revoke this release at any time in writing, but the Agency is not required to recall materials already published.")
    add_p(doc, "I understand that signing this release is OPTIONAL and is not a condition of receiving services.")

    add_signature_block(doc)
    return doc


def doc_emergency_contact() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "EMERGENCY CONTACT AND CRITICAL INFORMATION")

    add_p(doc, "Client Name: {{CLIENT_NAME}}    Date of Birth: {{CLIENT_DOB}}")
    add_p(doc, "Address: ______________________________________  Phone: __________________")

    add_h2(doc, "Primary Emergency Contact")
    add_p(doc, "Name: __________________________  Relationship: ______________")
    add_p(doc, "Phone (Home): _______________  (Cell): _______________  (Work): _______________")
    add_p(doc, "[ ] Has Power of Attorney   [ ] Has Healthcare Proxy")

    add_h2(doc, "Secondary Emergency Contact")
    add_p(doc, "Name: __________________________  Relationship: ______________")
    add_p(doc, "Phone (Home): _______________  (Cell): _______________  (Work): _______________")

    add_h2(doc, "Primary Care Physician")
    add_p(doc, "Name: __________________________  Practice: __________________________")
    add_p(doc, "Phone: __________________  Address: __________________________________")

    add_h2(doc, "Preferred Hospital")
    add_p(doc, "Hospital: ____________________________  City: ____________________________")

    add_h2(doc, "Allergies and Medical Alerts")
    add_p(doc, "________________________________________________________________________")
    add_p(doc, "________________________________________________________________________")

    add_h2(doc, "Current Medications (attach list if needed)")
    add_p(doc, "________________________________________________________________________")
    add_p(doc, "________________________________________________________________________")

    add_h2(doc, "Code Status / Advance Directive on File")
    add_p(doc, "[ ] Full Code   [ ] DNR (out-of-hospital order on file)   [ ] Other: __________")

    add_signature_block(doc)
    return doc


def doc_grievance_policy() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "GRIEVANCE AND COMPLAINT POLICY ACKNOWLEDGMENT")

    add_p(doc, "{{AGENCY_NAME}} encourages clients, family members, and authorized representatives to raise any concerns about services without fear of retaliation, reprisal, or discrimination.")

    add_h2(doc, "How to File a Grievance")
    add_p(doc, "Step 1: Speak directly with your assigned Care Coordinator, or call {{AGENCY_PHONE}}.")
    add_p(doc, "Step 2: If unresolved, submit your complaint in writing to {{COMPLIANCE_OFFICER_NAME}} at {{COMPLIANCE_OFFICER_EMAIL}} or by mail to {{AGENCY_ADDRESS}}, {{AGENCY_CITY}}, {{AGENCY_STATE}} {{AGENCY_ZIP}}.")
    add_p(doc, "Step 3: The Agency will acknowledge your complaint within {{GRIEVANCE_ACK_DAYS}} business days and provide a written response within {{GRIEVANCE_RESOLVE_DAYS}} business days.")

    add_h2(doc, "External Resources")
    add_bullets(doc, [
        "{{STATE_REGULATOR_NAME}} (state regulator): {{STATE_REGULATOR_PHONE}}",
        "Adult Protective Services: {{APS_PHONE}}",
        "Long-Term Care Ombudsman: {{OMBUDSMAN_PHONE}}",
        "U.S. Department of Health and Human Services Office for Civil Rights (HIPAA complaints): 1-800-368-1019",
    ])

    add_h2(doc, "Non-Retaliation")
    add_p(doc, "{{AGENCY_NAME}} prohibits retaliation against any person who files a complaint in good faith. Retaliation by any employee will result in disciplinary action up to and including termination.")

    add_signature_block(doc)
    return doc


# -----------------------------
# HIRING / HR DOCUMENTS
# -----------------------------


def doc_employment_application() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "EMPLOYMENT APPLICATION")
    add_subtitle(doc, "{{AGENCY_NAME}} is an Equal Opportunity Employer.")

    add_h2(doc, "Position Information")
    add_p(doc, "Position Applied For: ________________________________  Date: ____________")
    add_p(doc, "How did you hear about us? _____________________________________________")
    add_p(doc, "Desired Pay Rate: $__________  [ ] Hourly  [ ] Salary    Earliest Start Date: __________")
    add_p(doc, "Are you legally authorized to work in the United States? [ ] Yes  [ ] No")
    add_p(doc, "Will you require sponsorship now or in the future? [ ] Yes  [ ] No")

    add_h2(doc, "Personal Information")
    add_p(doc, "Full Legal Name: ____________________________________________________")
    add_p(doc, "Address: ____________________________________________________________")
    add_p(doc, "City: __________________________  State: ______  ZIP: __________")
    add_p(doc, "Phone: ____________________  Email: ____________________________")

    add_h2(doc, "Eligibility & Background")
    add_p(doc, "Have you ever been convicted of a felony or a crime involving violence, dishonesty, or abuse/neglect of a vulnerable person? [ ] Yes  [ ] No")
    add_p(doc, "If yes, please explain (a conviction does not automatically bar employment): ___________________________________________________________________")
    add_p(doc, "Do you have a valid driver's license? [ ] Yes  [ ] No   License #: __________  State: ______")
    add_p(doc, "Do you have reliable transportation? [ ] Yes  [ ] No")
    add_p(doc, "Do you have current auto insurance meeting state minimums? [ ] Yes  [ ] No")

    add_h2(doc, "Education")
    add_p(doc, "High School / GED: __________________________  Year: ______")
    add_p(doc, "College / Vocational: ________________________  Degree: ____________")
    add_p(doc, "Certifications (CNA, HHA, RN, LPN, CPR/First Aid, etc.): __________________________________")

    add_h2(doc, "Employment History (most recent first)")
    add_p(doc, "1. Employer: __________________  Position: __________  Dates: ______ to ______")
    add_p(doc, "   Reason for leaving: __________________  May we contact? [ ] Yes [ ] No")
    add_p(doc, "2. Employer: __________________  Position: __________  Dates: ______ to ______")
    add_p(doc, "   Reason for leaving: __________________  May we contact? [ ] Yes [ ] No")
    add_p(doc, "3. Employer: __________________  Position: __________  Dates: ______ to ______")
    add_p(doc, "   Reason for leaving: __________________  May we contact? [ ] Yes [ ] No")

    add_h2(doc, "References (non-relatives)")
    add_p(doc, "1. Name: ______________  Relationship: ____________  Phone: ______________")
    add_p(doc, "2. Name: ______________  Relationship: ____________  Phone: ______________")

    add_h2(doc, "Applicant Certification")
    add_p(doc, "I certify that the information provided in this application is true and complete to the best of my knowledge. I authorize {{AGENCY_NAME}} to verify the information, conduct background checks, and contact prior employers and references. I understand that any false statement or omission may result in withdrawal of any offer or termination of employment.")
    add_p(doc, "I understand that, if hired, my employment will be at-will (where applicable by law) and that no representation contrary to at-will employment is binding unless in writing signed by an officer of {{AGENCY_NAME}}.")

    doc.add_paragraph()
    doc.add_paragraph("Applicant Signature: ___________________________________  Date: ____________")
    return doc


def doc_offer_letter() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "OFFER OF EMPLOYMENT")
    add_p(doc, "{{OFFER_DATE}}")
    doc.add_paragraph()
    add_p(doc, "{{CANDIDATE_NAME}}")
    add_p(doc, "{{CANDIDATE_ADDRESS}}")
    doc.add_paragraph()
    add_p(doc, "Dear {{CANDIDATE_FIRST_NAME}},")
    add_p(doc, "On behalf of {{AGENCY_NAME}}, I am pleased to offer you the position of {{POSITION_TITLE}}, reporting to {{SUPERVISOR_NAME}}, with an anticipated start date of {{START_DATE}}.")

    add_h2(doc, "Compensation and Benefits")
    add_bullets(doc, [
        "Pay rate: ${{PAY_RATE}} per {{PAY_UNIT}} ({{HOURS_TYPE}})",
        "Pay schedule: {{PAY_SCHEDULE}} via direct deposit",
        "Eligibility for benefits: {{BENEFITS_SUMMARY}}",
        "Paid time off: {{PTO_SUMMARY}}",
        "Mileage reimbursement: ${{MILEAGE_RATE}}/mile for client-related travel",
    ])

    add_h2(doc, "Conditions of Employment")
    add_p(doc, "This offer is contingent upon successful completion of the following before your start date:")
    add_bullets(doc, [
        "Form I-9 employment eligibility verification",
        "Federal and state criminal background check",
        "{{AGENCY_STATE}}-required registry checks (e.g., abuse registries, OIG, sex offender)",
        "Driving record check (for positions that require driving)",
        "Pre-employment drug screen (if applicable)",
        "TB test or screening within the past 12 months",
        "Verification of all professional licenses, certifications, and credentials",
        "Signed Confidentiality, Code of Conduct, and Handbook acknowledgments",
    ])

    add_h2(doc, "At-Will Employment")
    add_p(doc, "Your employment with {{AGENCY_NAME}} is at-will, where permitted by law. This means either you or the Agency may terminate the employment relationship at any time, with or without notice or cause. No statement by anyone other than an officer of the Agency, in writing, can change the at-will nature of your employment.")

    add_h2(doc, "Acceptance")
    add_p(doc, "Please indicate your acceptance by signing below and returning this letter by {{OFFER_EXPIRATION_DATE}}. We are excited to have you join the team.")
    doc.add_paragraph()
    add_p(doc, "Sincerely,")
    add_p(doc, "{{HIRING_MANAGER_NAME}}, {{HIRING_MANAGER_TITLE}}")
    doc.add_paragraph()
    add_p(doc, "Accepted by: ___________________________________  Date: ____________")
    return doc


def doc_confidentiality_nda() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "CONFIDENTIALITY AND NON-DISCLOSURE AGREEMENT")

    add_p(doc, "This Agreement is between {{AGENCY_NAME}} (\"Agency\") and {{EMPLOYEE_NAME}} (\"Employee\") and is effective as of {{EFFECTIVE_DATE}}.")

    add_h2(doc, "1. Confidential Information")
    add_p(doc, "\"Confidential Information\" means any non-public information Employee receives or generates in connection with employment, including, without limitation: client and family identities; protected health information (PHI); plans of care; medication lists; financial and billing data; staffing schedules; pricing; trade secrets; software; passwords; vendor agreements; and personnel records.")

    add_h2(doc, "2. HIPAA Obligations")
    add_p(doc, "Employee agrees to comply with the Health Insurance Portability and Accountability Act of 1996 (HIPAA), the HITECH Act, and Agency's Notice of Privacy Practices and policies. Employee will only access PHI on a need-to-know basis to perform job duties and will not disclose PHI to unauthorized persons under any circumstance.")

    add_h2(doc, "3. Non-Use and Non-Disclosure")
    add_p(doc, "Employee will (a) hold all Confidential Information in strict confidence, (b) not use Confidential Information for any purpose other than performing duties for the Agency, and (c) not disclose Confidential Information to any third party without prior written authorization from the Agency.")

    add_h2(doc, "4. No Photography or Recording")
    add_p(doc, "Employee will not photograph, video, or audio-record clients, family members, the home environment, or Agency materials without explicit written authorization from the Agency and, where required, the client.")

    add_h2(doc, "5. Social Media and Public Communications")
    add_p(doc, "Employee will not post any client information, photos, complaints, or details about Agency operations on social media or any public forum. Personal posts that identify or could reasonably identify a client are strictly prohibited.")

    add_h2(doc, "6. Return of Materials")
    add_p(doc, "Upon termination of employment for any reason, Employee will immediately return all Agency property, including documents, devices, keys, uniforms, ID badges, and any copies of Confidential Information.")

    add_h2(doc, "7. Survival")
    add_p(doc, "Employee's obligations under this Agreement survive termination of employment indefinitely as to trade secrets and PHI, and for {{NDA_SURVIVAL_YEARS}} years as to other Confidential Information.")

    add_h2(doc, "8. Remedies")
    add_p(doc, "Employee acknowledges that breach of this Agreement may cause irreparable harm and that the Agency is entitled to seek injunctive relief in addition to any other remedies, including under HIPAA, state privacy laws, and trade-secret statutes.")

    add_h2(doc, "9. Governing Law")
    add_p(doc, "This Agreement is governed by the laws of the State of {{AGENCY_STATE}}.")

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Printed Name: ______________________________________________")
    doc.add_paragraph()
    add_p(doc, "Agency Representative: ________________________________  Date: ____________")
    return doc


def doc_background_check_auth() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "BACKGROUND CHECK DISCLOSURE AND AUTHORIZATION")
    add_subtitle(doc, "Fair Credit Reporting Act (FCRA) — Standalone Disclosure")

    add_h2(doc, "Disclosure")
    add_p(doc, "{{AGENCY_NAME}} (\"Agency\") may obtain a consumer report and/or investigative consumer report about you for employment purposes, including hiring, promotion, reassignment, or retention. The report may include information about your character, general reputation, personal characteristics, mode of living, criminal history, motor vehicle records, employment history, education, professional licenses, and credentials.")
    add_p(doc, "Reports will be obtained from a consumer reporting agency. You may request additional information about the nature and scope of the report.")

    add_h2(doc, "Healthcare-Specific Checks")
    add_p(doc, "Because the Agency provides care to vulnerable populations, the report may also include:")
    add_bullets(doc, [
        "Federal criminal history (FBI fingerprint-based check)",
        "State criminal history (per {{AGENCY_STATE}} law)",
        "OIG List of Excluded Individuals/Entities (LEIE)",
        "GSA System for Award Management (SAM) exclusions",
        "State Medicaid exclusion lists",
        "State abuse, neglect, and elder/child protective services registries",
        "National Sex Offender Public Website",
        "Professional license verification and disciplinary history",
    ])

    add_h2(doc, "Authorization")
    add_p(doc, "I have read and understand the above Disclosure. I authorize {{AGENCY_NAME}} and its consumer reporting agency to obtain the report described and any updates from time to time during my employment. I authorize all parties holding such information — including former employers, schools, government agencies, and licensing boards — to release the information requested.")

    add_h2(doc, "Applicant Information (printed)")
    add_p(doc, "Full Legal Name: ____________________________________________________")
    add_p(doc, "Other Names Used (maiden, alias): ___________________________________")
    add_p(doc, "Date of Birth: __________  Last 4 of SSN: ________  Driver's License #: __________  State: ____")
    add_p(doc, "Current Address: ____________________________________________________")
    add_p(doc, "Prior Addresses (last 7 years): _____________________________________")

    doc.add_paragraph()
    add_p(doc, "Applicant Signature: ___________________________________  Date: ____________")

    add_h2(doc, "California, Minnesota, Oklahoma Applicants")
    add_p(doc, "[ ] I would like a copy of any consumer report obtained about me.")
    return doc


def doc_drug_testing_consent() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "DRUG AND ALCOHOL TESTING CONSENT")

    add_p(doc, "{{AGENCY_NAME}} maintains a drug- and alcohol-free workplace to protect the health and safety of clients, caregivers, and the public.")

    add_h2(doc, "Testing Conditions")
    add_bullets(doc, [
        "Pre-employment testing as a condition of hire",
        "Reasonable-suspicion testing based on observed behavior or performance",
        "Post-incident testing following any workplace incident or injury",
        "Random testing where permitted by state law",
        "Return-to-duty and follow-up testing after a positive result or treatment",
    ])

    add_h2(doc, "Substances Tested")
    add_p(doc, "Testing may screen for amphetamines, cocaine, opiates, phencyclidine (PCP), THC/marijuana metabolites, benzodiazepines, barbiturates, methadone, propoxyphene, and alcohol, in accordance with {{AGENCY_STATE}} law.")

    add_h2(doc, "Authorization")
    add_p(doc, "I voluntarily consent to drug and/or alcohol testing as described. I authorize the testing facility and the Agency's Medical Review Officer to release the results to {{AGENCY_NAME}}. I understand that refusal to test, tampering with a sample, or a confirmed positive test may result in withdrawal of any offer or termination of employment, subject to applicable law.")

    add_h2(doc, "Prescription Medications")
    add_p(doc, "If a positive test is reported, the Medical Review Officer will contact me to discuss any prescription or over-the-counter medications that may explain the result. I am not required to disclose prescriptions in advance.")

    doc.add_paragraph()
    add_p(doc, "Employee/Applicant Signature: ___________________________________  Date: ____________")
    return doc


def doc_direct_deposit() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "DIRECT DEPOSIT AUTHORIZATION")

    add_p(doc, "I authorize {{AGENCY_NAME}} to deposit my net pay (and any reimbursements) directly into the account(s) listed below. This authorization remains in effect until I revoke it in writing.")

    add_h2(doc, "Account 1 (Primary)")
    add_p(doc, "Bank Name: ______________________________________________________")
    add_p(doc, "Routing # (9 digits): ____________  Account #: __________________")
    add_p(doc, "[ ] Checking   [ ] Savings    Amount: [ ] 100% of net pay   [ ] $__________")

    add_h2(doc, "Account 2 (Optional)")
    add_p(doc, "Bank Name: ______________________________________________________")
    add_p(doc, "Routing # (9 digits): ____________  Account #: __________________")
    add_p(doc, "[ ] Checking   [ ] Savings    Amount: [ ] Remainder   [ ] $__________")

    add_h2(doc, "Verification")
    add_p(doc, "Please attach a voided check or a direct deposit verification letter from your bank for each account above.")

    add_h2(doc, "Authorization")
    add_p(doc, "I authorize {{AGENCY_NAME}} to credit and, if necessary to correct an error, debit the account(s) above. I understand any changes require a new signed authorization and may take up to one pay cycle to take effect.")

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Printed Name: ______________________________________________")
    add_p(doc, "Last 4 of SSN: __________  Employee ID: ______________")
    return doc


def doc_handbook_ack() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "EMPLOYEE HANDBOOK ACKNOWLEDGMENT")

    add_p(doc, "I acknowledge that I have received a copy of the {{AGENCY_NAME}} Employee Handbook (\"Handbook\") dated {{HANDBOOK_DATE}}. I understand that:")
    add_bullets(doc, [
        "The Handbook contains important information about Agency policies, procedures, benefits, and expectations.",
        "I am responsible for reading the Handbook and following all policies it contains.",
        "Policies in the Handbook may change at any time, with or without notice, at the sole discretion of the Agency.",
        "The Handbook is not a contract of employment and does not create any guarantee of employment for any duration.",
        "My employment is at-will, where permitted by law. Either the Agency or I may end the employment relationship at any time, with or without notice or cause.",
        "Only an officer of the Agency may modify at-will status, and only in writing.",
        "I am responsible for asking my supervisor or HR for clarification on any policy I do not understand.",
    ])

    add_h2(doc, "Specific Policies Acknowledged")
    add_bullets(doc, [
        "Equal Employment Opportunity and Anti-Harassment",
        "HIPAA / Confidentiality",
        "Drug- and Alcohol-Free Workplace",
        "Code of Conduct and Ethics",
        "Mandatory Reporting (suspected abuse, neglect, exploitation)",
        "Electronic Communications and Social Media",
        "Time Reporting and EVV Compliance",
        "Attendance and Punctuality",
        "Workplace Safety and Bloodborne Pathogens",
        "Conflict of Interest and Acceptance of Gifts",
        "Discipline and Termination Procedures",
    ])

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Printed Name: ______________________________________________")
    return doc


def doc_code_of_conduct() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "CODE OF CONDUCT AND ETHICS ACKNOWLEDGMENT")

    add_p(doc, "All caregivers, employees, and contractors of {{AGENCY_NAME}} are expected to uphold the highest standards of professional and ethical conduct. By signing below, I agree to the following standards:")

    add_h2(doc, "Professional Conduct")
    add_bullets(doc, [
        "Treat every client, family member, coworker, and member of the public with respect, dignity, and compassion.",
        "Maintain professional boundaries; do not engage in romantic, sexual, or financial relationships with clients or family members.",
        "Do not accept cash, gifts of significant value, loans, tips, bequests, or property from clients or family members.",
        "Do not borrow from, lend to, or co-mingle finances with clients.",
        "Do not transport clients in personal vehicles unless explicitly authorized in writing.",
    ])

    add_h2(doc, "Ethical Conduct")
    add_bullets(doc, [
        "Document care accurately and contemporaneously. Falsifying records, time entries, or EVV check-ins is grounds for immediate termination and may result in criminal charges.",
        "Bill only for services actually rendered. Submitting false claims to Medicare, Medicaid, or insurers may result in criminal and civil liability under the False Claims Act.",
        "Disclose any conflicts of interest, including outside employment with the same client or family.",
        "Comply with HIPAA, state privacy laws, and Agency policies regarding PHI.",
    ])

    add_h2(doc, "Mandatory Reporting")
    add_bullets(doc, [
        "Report any suspected abuse, neglect, exploitation, or self-neglect of any client immediately to a supervisor and to Adult Protective Services or the appropriate state agency.",
        "Report any workplace injury or client incident immediately.",
        "Report any suspected fraud, theft, or unethical conduct via the Agency's confidential compliance line at {{COMPLIANCE_HOTLINE}}.",
    ])

    add_h2(doc, "Discrimination and Harassment")
    add_p(doc, "{{AGENCY_NAME}} prohibits discrimination and harassment based on race, color, religion, national origin, sex, sexual orientation, gender identity, age, disability, marital status, pregnancy, veteran status, or any other protected category. This applies to coworkers, clients, family members, and visitors.")

    add_h2(doc, "Consequences")
    add_p(doc, "Violation of this Code may result in disciplinary action up to and including termination, criminal prosecution, and civil liability.")

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Printed Name: ______________________________________________")
    return doc


def doc_hipaa_training_ack() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "HIPAA TRAINING ACKNOWLEDGMENT")

    add_p(doc, "I acknowledge that I have completed HIPAA Privacy and Security training provided by {{AGENCY_NAME}} on {{TRAINING_DATE}}, including the following topics:")
    add_bullets(doc, [
        "What is Protected Health Information (PHI) and what is not",
        "Permissible uses and disclosures (Treatment, Payment, Operations)",
        "Minimum-necessary standard",
        "Patient rights under HIPAA",
        "Safeguards for paper, electronic, and verbal PHI",
        "Use of personal devices, email, and text messaging",
        "Social media restrictions",
        "Breach identification and reporting (within 24 hours to Privacy Officer)",
        "Sanctions for HIPAA violations, including civil and criminal penalties",
    ])

    add_p(doc, "I understand that I must complete annual HIPAA refresher training. I understand that any unauthorized access, use, or disclosure of PHI may result in disciplinary action up to and including termination, civil and criminal penalties under HIPAA (up to $1.5 million per violation category and 10 years' imprisonment for knowing willful disclosures), and personal liability under state law.")

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Trainer Signature: ____________________________________  Date: ____________")
    return doc


def doc_mandatory_reporter() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "MANDATORY REPORTER ACKNOWLEDGMENT")

    add_p(doc, "Under the laws of {{AGENCY_STATE}} and applicable federal law, employees and contractors of {{AGENCY_NAME}} who provide care to vulnerable adults and/or children are mandated reporters of suspected abuse, neglect, exploitation, or self-neglect.")

    add_h2(doc, "What to Report")
    add_bullets(doc, [
        "Physical abuse or signs of physical injury (bruises, burns, fractures)",
        "Sexual abuse or sexually inappropriate contact",
        "Emotional or psychological abuse",
        "Neglect, including caregiver neglect, self-neglect, and abandonment",
        "Financial exploitation, theft, or coercion",
        "Unsafe or unsanitary living conditions that endanger the client",
    ])

    add_h2(doc, "How and When to Report")
    add_bullets(doc, [
        "Report immediately — do not wait until the end of your shift.",
        "First, ensure the client is safe and call 911 if there is an immediate threat.",
        "Report to your supervisor at {{AGENCY_PHONE}}.",
        "Report to Adult Protective Services / Child Protective Services as required by law: {{APS_PHONE}} / {{CPS_PHONE}}.",
        "Document what you observed (objective facts, not opinions) as soon as practical.",
    ])

    add_h2(doc, "Confidentiality and Protection")
    add_p(doc, "Reports are confidential to the extent permitted by law. Mandated reporters acting in good faith are immune from civil and criminal liability. Retaliation against a mandated reporter is prohibited and is itself a reportable offense.")

    add_h2(doc, "Penalties for Failure to Report")
    add_p(doc, "Failure to report suspected abuse, neglect, or exploitation may result in criminal penalties under {{AGENCY_STATE}} law, civil liability, loss of professional licensure, and termination.")

    doc.add_paragraph()
    add_p(doc, "Employee Signature: ___________________________________  Date: ____________")
    add_p(doc, "Printed Name: ______________________________________________")
    return doc


def doc_caregiver_job_description() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "CAREGIVER / HOME HEALTH AIDE — JOB DESCRIPTION")

    add_p(doc, "Position Title: {{POSITION_TITLE}}    Reports To: {{SUPERVISOR_TITLE}}")
    add_p(doc, "FLSA Status: {{FLSA_STATUS}}    Pay Type: {{PAY_UNIT}}")

    add_h2(doc, "Position Summary")
    add_p(doc, "The Caregiver provides non-medical and/or skilled in-home support services to clients of {{AGENCY_NAME}} in accordance with the client's individualized Plan of Care. The Caregiver promotes client safety, dignity, and independence in their home environment.")

    add_h2(doc, "Essential Duties")
    add_bullets(doc, [
        "Personal care: bathing, grooming, dressing, toileting, incontinence care, transfers, and ambulation",
        "Meal planning and preparation per dietary needs",
        "Light housekeeping, laundry, and bed-making",
        "Medication reminders within scope of practice",
        "Companionship, supervision, and engagement in client-preferred activities",
        "Errands, transportation, and accompaniment to appointments (when authorized)",
        "Accurate, contemporaneous documentation in the Agency's EVV/charting system",
        "Observation and reporting of changes in client condition to the supervisor or nurse",
        "Compliance with HIPAA, infection control, and Agency policies at all times",
    ])

    add_h2(doc, "Required Qualifications")
    add_bullets(doc, [
        "High school diploma or GED",
        "Active state-required certification (CNA, HHA, or equivalent) where applicable",
        "Current CPR / First Aid certification",
        "Negative TB screening within the past 12 months",
        "Pass federal and state background checks, registry checks, and drug screen",
        "Valid driver's license, reliable transportation, and current auto insurance (for positions requiring driving)",
        "Ability to read, write, and communicate in English at a level sufficient for safe care",
        "Ability to follow written and verbal instructions",
    ])

    add_h2(doc, "Physical Requirements")
    add_bullets(doc, [
        "Lift up to 50 pounds and assist with transfers; may use mechanical lifts",
        "Stand, walk, bend, kneel, and reach for extended periods",
        "Push, pull, and operate household equipment",
        "Tolerate exposure to body fluids, infectious materials, pets, and household allergens",
    ])

    add_h2(doc, "Working Conditions")
    add_bullets(doc, [
        "Work is performed in clients' private homes, which may vary in cleanliness, safety, and accessibility.",
        "Schedule may include weekends, evenings, overnights, and holidays.",
        "Travel between client homes is required for some positions.",
    ])

    add_p(doc, "This job description is not exhaustive. The Agency reserves the right to add or change duties as business needs require.")
    doc.add_paragraph()
    add_p(doc, "Acknowledged by: ___________________________________  Date: ____________")
    return doc


def doc_i9_w4_coversheet() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "I-9 / W-4 / NEW HIRE PAPERWORK COVER SHEET")
    add_subtitle(doc, "Federal forms must be the most current version, downloaded directly from the issuing agency.")

    add_h2(doc, "Federal Forms (Required)")
    add_bullets(doc, [
        "Form I-9 — Employment Eligibility Verification (uscis.gov/i-9). Section 1 completed by employee on or before first day of work; Section 2 completed by Agency within 3 business days.",
        "Form W-4 — Employee's Withholding Certificate (irs.gov/formspubs).",
        "(If applicable) Form W-4P, Form 8233 (treaty nonresidents), or Form W-9 (independent contractors).",
    ])

    add_h2(doc, "State Forms (varies by state)")
    add_bullets(doc, [
        "{{AGENCY_STATE}} state withholding form (if applicable)",
        "New-hire reporting form (the Agency files within 20 days of hire)",
        "Workers' compensation acknowledgment",
        "Unemployment insurance information",
    ])

    add_h2(doc, "I-9 Acceptable Documents — Common Combinations")
    add_p(doc, "List A (one document — establishes both identity and employment authorization): U.S. Passport, Permanent Resident Card (Green Card), Foreign Passport with I-551 stamp, Employment Authorization Document (EAD).")
    add_p(doc, "OR")
    add_p(doc, "List B (one — identity) + List C (one — employment authorization). Examples: Driver's license + Social Security card; State ID + birth certificate.")

    add_h2(doc, "Agency Verification")
    add_p(doc, "The Agency verifies and copies acceptable I-9 documents. {{AGENCY_NAME}} {{E_VERIFY_STATUS}} an E-Verify employer.")
    add_p(doc, "I-9 records are retained for {{I9_RETENTION_TEXT}}, separately from personnel files.")

    doc.add_paragraph()
    add_p(doc, "Employee acknowledgment of receipt: __________________________  Date: ____________")
    return doc


# -----------------------------
# STATE ADDENDA — PA & OH
# -----------------------------


def doc_pa_addendum() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "PENNSYLVANIA STATE ADDENDUM")
    add_subtitle(doc, "This addendum supplements the Service Agreement and Hiring Documents for operations in the Commonwealth of Pennsylvania.")
    add_p(doc, "INFORMATIONAL — review with Pennsylvania-licensed counsel before relying on this content. Citations reflect law as of the document date and may change.")

    add_h2(doc, "1. Pennsylvania Home Care Agency Licensure")
    add_p(doc, "{{AGENCY_NAME}} is licensed by the Pennsylvania Department of Health under Act 69 of 2006 / Chapter 611 of Title 28 (Home Care Agencies and Home Care Registries). License #: {{PA_HOMECARE_LICENSE}}. Complaints may be filed with the PA DOH Division of Home Health at 1-800-254-5164.")

    add_h2(doc, "2. Background Checks (Caregiver Hiring)")
    add_p(doc, "Pennsylvania requires the following pre-employment screenings before a direct care worker may be hired or assigned to a client:")
    add_bullets(doc, [
        "Pennsylvania State Police Criminal Record Check (Act 153 / Act 73 / Act 169) — required for all employees and contractors with direct client contact.",
        "Pennsylvania Child Abuse History Clearance (ChildLine — Act 153) where any client is under 18 or where caregivers may have contact with minors.",
        "FBI Federal Criminal Background Check (fingerprint-based) for any applicant who has not been a Pennsylvania resident continuously for the past two years.",
        "Older Adults Protective Services Act (OAPSA, 35 P.S. §10225.502) certifications, including review of disqualifying offenses.",
        "Department of Aging registry checks where applicable.",
    ])
    add_p(doc, "Any positive criminal history will be evaluated under OAPSA disqualification rules. Agency will follow Act 13 of 2017 procedures, including individualized assessment where applicable.")

    add_h2(doc, "3. Mandatory Reporting (PA)")
    add_p(doc, "Caregivers and Agency staff are mandatory reporters of suspected abuse, neglect, exploitation, or abandonment of older adults under OAPSA, and of suspected child abuse under the Child Protective Services Law (CPSL, 23 Pa.C.S. § 6311).")
    add_bullets(doc, [
        "Older adults: PA Department of Aging Elder Abuse Hotline 1-800-490-8505",
        "Child abuse: PA ChildLine 1-800-932-0313",
        "Reports must be made immediately and followed up in writing within 48 hours where required.",
    ])

    add_h2(doc, "4. Caregiver Training Requirements")
    add_p(doc, "Per 28 Pa. Code § 611.55, direct care workers must complete competency requirements prior to providing care, including: client rights, infection control, basic safety, recognition of abuse/neglect, confidentiality, and CPR/First Aid (where required by job duties). Annual continuing education is required.")

    add_h2(doc, "5. Wage and Hour")
    add_p(doc, "Pennsylvania minimum wage and the Pennsylvania Minimum Wage Act apply. As of {{EFFECTIVE_DATE}}, PA's Final Rule on Overtime (eff. 2022) and the federal home-care companionship rule (29 C.F.R. § 552) apply to most direct-care workers, who are entitled to minimum wage and overtime for hours over 40 per week. Agency will track all hours worked, including travel between client homes when compensable.")

    add_h2(doc, "6. Workers' Compensation")
    add_p(doc, "All employees are covered by Pennsylvania workers' compensation insurance carried by {{AGENCY_NAME}}. Carrier: {{WC_CARRIER}}. Report any work-related injury to your supervisor and to {{WC_CONTACT}} within 24 hours.")

    add_h2(doc, "7. Client Rights (PA-specific)")
    add_p(doc, "In addition to the rights listed in the Bill of Rights, Pennsylvania consumers have the right under 28 Pa. Code § 611.57 to receive written notice of services, fees, and the consumer's right to file a complaint with the PA Department of Health.")

    add_signature_block(doc)
    return doc


def doc_oh_addendum() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "OHIO STATE ADDENDUM")
    add_subtitle(doc, "This addendum supplements the Service Agreement and Hiring Documents for operations in the State of Ohio.")
    add_p(doc, "INFORMATIONAL — review with Ohio-licensed counsel before relying on this content. Citations reflect law as of the document date and may change.")

    add_h2(doc, "1. Ohio Licensure / Certification")
    add_p(doc, "Depending on services offered, {{AGENCY_NAME}} operates under one or more of the following: Ohio Department of Health Home Health Agency certification (Medicare/Medicaid), Ohio Department of Aging PASSPORT/MyCare provider certification (ORC Chapter 173), and/or Ohio Department of Medicaid provider agreement under the Home Care Waiver (OAC 5160). License/Provider #: {{OH_PROVIDER_NUMBER}}.")

    add_h2(doc, "2. Background Checks (Caregiver Hiring)")
    add_p(doc, "Ohio requires the following pre-employment screenings before a direct care worker may be hired or assigned to a client:")
    add_bullets(doc, [
        "Ohio Bureau of Criminal Investigation (BCI) state criminal records check",
        "FBI federal criminal records check for any applicant who has not been an Ohio resident continuously for the past five years",
        "Database checks per ORC § 109.572 and OAC 3701-13-06 / 5160-45-11, including: OIG List of Excluded Individuals/Entities, GSA System for Award Management, U.S. Treasury OFAC, Ohio Sex Offender registry, Ohio Nurse Aide Registry abuse findings, and the Department of Developmental Disabilities Abuser Registry.",
        "Driver's license abstract from the Ohio Bureau of Motor Vehicles for any caregiver who will transport clients.",
    ])
    add_p(doc, "Disqualifying offenses are listed in ORC § 109.572. Re-checks are performed every five (5) years per OAC 3701-13.")

    add_h2(doc, "3. Mandatory Reporting (OH)")
    add_p(doc, "Under ORC § 5101.61, certain employees and providers are mandated reporters of adult abuse, neglect, and exploitation. Under ORC § 2151.421, certain employees are mandated reporters of child abuse and neglect.")
    add_bullets(doc, [
        "Ohio Adult Protective Services (county-based): 1-855-OHIO-APS (1-855-644-6277)",
        "Ohio Child Protective Services / 24-hour hotline: 1-855-O-H-CHILD (1-855-642-4453)",
        "Ohio Department of Health complaint hotline: 1-800-342-0553",
        "Reports must be made immediately upon suspicion.",
    ])

    add_h2(doc, "4. Training and Competency")
    add_p(doc, "Direct care workers must satisfy the training and competency requirements of OAC 3701-17 (home health aides) and/or OAC 5160-45-08 (waiver personal care aides), including initial competency evaluation, infection control, HIPAA, recognition of abuse/neglect, CPR/First Aid where required, and 12 hours of in-service training annually.")

    add_h2(doc, "5. Wage and Hour")
    add_p(doc, "Ohio minimum wage applies. Federal FLSA companionship rules (29 C.F.R. § 552) apply to most direct-care workers, who are entitled to minimum wage and overtime for hours over 40 per week. Ohio prohibits discharge in retaliation for filing a wage claim (ORC § 4111.13).")

    add_h2(doc, "6. Workers' Compensation")
    add_p(doc, "All Ohio employees are covered by the Ohio Bureau of Workers' Compensation (BWC) state fund. Policy #: {{OH_BWC_POLICY}}. Report any work-related injury to your supervisor and to BWC within 24 hours.")

    add_h2(doc, "7. Client Rights (OH-specific)")
    add_p(doc, "In addition to the Bill of Rights, Ohio consumers have rights under ORC § 3721.13 / OAC 3701-17-09, including the right to receive written notice of services, fees, and the right to file a complaint with the Ohio Department of Health (1-800-342-0553), the Long-Term Care Ombudsman (1-800-282-1206), or, for waiver participants, the Ohio Department of Medicaid.")

    add_signature_block(doc)
    return doc


# -----------------------------
# DRIVER ADDENDUM (HIRING)
# -----------------------------


def doc_driver_addendum() -> Document:
    doc = new_doc()
    agency_header(doc)
    add_title(doc, "DRIVER AND VEHICLE VERIFICATION")

    add_p(doc, "Required only for caregivers who will drive themselves or transport clients in any vehicle in connection with their duties.")

    add_h2(doc, "Driver Information")
    add_p(doc, "Full Legal Name: ____________________________________________________")
    add_p(doc, "Driver's License #: ____________  State: ____  Expiration: __________")
    add_p(doc, "Any restrictions on license? [ ] No   [ ] Yes (specify): __________________________")
    add_p(doc, "Have you had your license suspended or revoked in the past 7 years? [ ] No   [ ] Yes")
    add_p(doc, "Have you been convicted of any moving violation, DUI/OVI, or reckless operation in the past 7 years? [ ] No   [ ] Yes (explain): __________________________")

    add_h2(doc, "Vehicle Information (if using personal vehicle)")
    add_p(doc, "Make / Model / Year: __________________________  Color: __________  Plate #: __________")
    add_p(doc, "Vehicle is currently registered and inspected (where required): [ ] Yes  [ ] No")

    add_h2(doc, "Auto Insurance")
    add_p(doc, "Carrier: ______________________________  Policy #: ______________________________")
    add_p(doc, "Liability limits: BI $______________ per person / $______________ per accident; PD $______________")
    add_p(doc, "Limits meet or exceed {{AGENCY_STATE}} statutory minimums and the Agency-required minimums of $100,000 / $300,000 / $50,000.")

    add_h2(doc, "Authorization")
    add_p(doc, "I authorize {{AGENCY_NAME}} to obtain my Motor Vehicle Record (MVR) at hire and at any time during employment. I will notify the Agency within 24 hours of any change in license status, insurance status, or moving violation. I understand that loss of a valid license, insurance, or an unsatisfactory MVR may result in reassignment or termination.")

    doc.add_paragraph()
    add_p(doc, "Driver Signature: ___________________________________  Date: ____________")
    return doc


# -----------------------------
# RUNNER
# -----------------------------


CLIENT_DOCS = [
    ("01-service-agreement.docx", doc_service_agreement),
    ("02-hipaa-privacy-notice.docx", doc_hipaa_notice),
    ("03-hipaa-authorization.docx", doc_hipaa_authorization),
    ("04-bill-of-rights.docx", doc_bill_of_rights),
    ("05-consent-to-care.docx", doc_consent_to_care),
    ("06-financial-responsibility.docx", doc_financial_responsibility),
    ("07-advance-directive-acknowledgment.docx", doc_advance_directive_ack),
    ("08-photo-media-release.docx", doc_photo_media_release),
    ("09-emergency-contact.docx", doc_emergency_contact),
    ("10-grievance-policy.docx", doc_grievance_policy),
]

HIRING_DOCS = [
    ("01-employment-application.docx", doc_employment_application),
    ("02-offer-letter.docx", doc_offer_letter),
    ("03-confidentiality-nda.docx", doc_confidentiality_nda),
    ("04-background-check-authorization.docx", doc_background_check_auth),
    ("05-drug-testing-consent.docx", doc_drug_testing_consent),
    ("06-direct-deposit-authorization.docx", doc_direct_deposit),
    ("07-employee-handbook-acknowledgment.docx", doc_handbook_ack),
    ("08-code-of-conduct.docx", doc_code_of_conduct),
    ("09-hipaa-training-acknowledgment.docx", doc_hipaa_training_ack),
    ("10-mandatory-reporter-acknowledgment.docx", doc_mandatory_reporter),
    ("11-caregiver-job-description.docx", doc_caregiver_job_description),
    ("12-i9-w4-coversheet.docx", doc_i9_w4_coversheet),
    ("13-driver-vehicle-verification.docx", doc_driver_addendum),
]

STATE_DOCS = [
    ("pennsylvania-addendum.docx", doc_pa_addendum),
    ("ohio-addendum.docx", doc_oh_addendum),
]


def main() -> None:
    print("Generating client onboarding documents:")
    for fname, builder in CLIENT_DOCS:
        save(builder(), CLIENT_DIR / fname)

    print("\nGenerating hiring documents:")
    for fname, builder in HIRING_DOCS:
        save(builder(), HIRING_DIR / fname)

    print("\nGenerating state addenda:")
    for fname, builder in STATE_DOCS:
        save(builder(), STATE_DIR / fname)

    print("\nDone.")


if __name__ == "__main__":
    main()
